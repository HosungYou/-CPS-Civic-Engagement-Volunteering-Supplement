from pathlib import Path

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml.ns import qn


SRC = Path("/Users/hosung/Downloads/Bowling_Alone_Scrolling_Together_Manuscript_Draft SW Feedback 03_11_2026.docx")


def set_paragraph_text(paragraph, text, highlight=False):
    paragraph.text = ""
    run = paragraph.add_run(text)
    if highlight:
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    return run


def clear_numbering(paragraph):
    p_pr = paragraph._p.pPr
    if p_pr is None:
        return
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is not None:
        p_pr.remove(num_pr)


def add_paragraph_before(before_paragraph, text, style=None, highlight=False):
    p = before_paragraph.insert_paragraph_before()
    if style:
        p.style = style
    clear_numbering(p)
    run = p.add_run(text)
    if highlight:
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    return p


def insert_table_before(paragraph, rows, cols):
    table = paragraph.part.document.add_table(rows=rows, cols=cols)
    paragraph._p.addprevious(table._tbl)
    return table


def revise_document(output_path: Path):
    doc = Document(str(SRC))
    paragraphs = list(doc.paragraphs)

    revisions = {
        12: (
            "Social isolation has increasingly been framed as a public health concern, yet its implications for volunteering remain underexamined. Using four waves of the Current Population Survey Civic Engagement and Volunteering Supplement (2017, 2019, 2021, and 2023; N = 201,168), this study examines whether in-person socialization is associated with formal volunteering across generational cohorts. The study adopts a sequential multi-method design that combines survey-weighted logistic regression, person-centered profile analysis, and gradient boosting with TreeSHAP to assess both theory-driven associations and heterogeneous engagement patterns. Particular attention is given to Generation Z, whose comparatively low levels of face-to-face socialization raise questions about whether traditional pathways to volunteering operate similarly across cohorts. Education, employment, and civic social media use are examined as moderators of the socialization-volunteering relationship. The study clarifies how changing patterns of social connection may shape volunteer recruitment and community participation."
        ),
        17: (
            "In May 2023, the U.S. Surgeon General described loneliness and social isolation as a major public health concern, emphasizing consequences for mental and physical well-being as well as community resilience (Office of the Surgeon General, 2023). Although the advisory noted that socially connected communities tend to show stronger civic engagement, the empirical discussion centered overwhelmingly on health outcomes. Less attention has been given to how declining face-to-face connection may affect organized volunteering, which remains one of the most visible ways individuals contribute time and labor to community life."
        ),
        18: (
            "This gap matters because volunteering depends heavily on social recruitment. Unlike lower-cost forms of participation such as signing petitions or making donations, formal volunteering usually requires repeated interaction with organizations, exposure to recruitment appeals, and enough social energy to sustain involvement over time (Musick & Wilson, 2008; Wilson, 2000). If social isolation reduces access to those recruitment pathways, then the consequences extend beyond individual nonparticipation to the staffing capacity of nonprofits and community organizations."
        ),
        19: (
            "The question is especially salient for younger cohorts. Recent scholarship documents declines in face-to-face socialization among adolescents and emerging adults, particularly after the spread of smartphones and digitally mediated interaction (Twenge, 2017; Twenge et al., 2019). At the same time, scholars of civic engagement note that younger cohorts may substitute personalized or digitally mediated action for organization-based participation (Bennett & Segerberg, 2012; Zukin et al., 2006). What remains unclear is whether lower levels of in-person socialization help explain cross-generational differences in formal volunteering and whether conventional civic resources operate similarly for Generation Z and older cohorts."
        ),
        20: (
            "Using nationally representative CPS-CEV data from 2017, 2019, 2021, and 2023, this study examines whether in-person socialization is associated with formal volunteering across generations and whether education, employment, and civic social media use condition that association. The analysis is organized to clarify both the substantive relationship between social connection and volunteering and the methodological value of combining variable-centered, person-centered, and machine-learning approaches."
        ),
        21: (
            "Rather than assuming that all younger adults disengage for the same reasons, the manuscript treats Generation Z as an analytically important case for testing whether diminished face-to-face interaction constrains access to volunteering opportunities."
        ),
        22: (
            "It also distinguishes between organized volunteering and broader forms of civic activity so that claims remain aligned with the study's dependent variable rather than with civic engagement in the abstract."
        ),
        23: (
            "The sections that follow review the literature on in-person socialization, volunteering across generations, and the potential moderating roles of education, employment, and civic social media use."
        ),
        24: "Literature Review",
        25: "In-Person Socialization and Volunteering",
        26: (
            "Social capital theory and the civic voluntarism model both suggest that in-person socialization should matter for volunteering. Putnam (2000) conceptualized social capital as the networks and norms that facilitate reciprocity, while Verba et al. (1995) emphasized that civic participation depends not only on resources and engagement but also on recruitment. For volunteering in particular, recruitment often occurs through routine interactions with friends, relatives, neighbors, coworkers, and fellow members of organizations (Brady et al., 1995; Musick & Wilson, 2008)."
        ),
        27: (
            "This perspective implies that the movement from no social contact to even occasional social contact may be substantively important. Coleman (1988) argued that social relationships create information channels and obligations that facilitate collective action, and Granovetter's (1973) theory of weak ties suggests that even casual acquaintances can connect individuals to new opportunities. Applied to volunteering, these frameworks support the expectation that more frequent in-person socialization will be associated with a higher likelihood of formal volunteering and that the steepest increase may occur at the lowest end of the socialization distribution."
        ),
        28: "Generational Differences in Volunteering: Why Generation Z?",
        29: (
            "Any account of socialization and volunteering must also address generational change. Twenge et al. (2019) documented substantial declines in face-to-face interaction among U.S. adolescents beginning in the early 2010s, while Twenge (2017) argued that the cohort born after the mid-1990s came of age in an environment where social life is more heavily mediated by digital technologies. These shifts make Generation Z a critical case for asking whether reduced in-person interaction is linked to lower volunteering."
        ),
        30: (
            "Scholars of civic engagement likewise show that participation has changed across cohorts, but not always in the same direction or through the same mechanisms. Younger adults may engage in consumer activism, issue expression, or episodic participation even as membership-based organizations lose centrality (Bennett & Segerberg, 2012; Dalton, 2008; Skocpol, 2003). That literature is useful, but it does not resolve whether formal volunteering follows the same trajectory. Because volunteering depends on organizational entry and retention, Generation Z may face a distinct disadvantage if reduced face-to-face socialization weakens access to recruitment networks."
        ),
        31: (
            "Generation is therefore treated in this study not as a stand-alone explanation but as a contextual moderator that may alter how strongly in-person socialization translates into volunteering. This framing keeps the analysis focused on a specific mechanism rather than on broad claims about the civic character of younger adults."
        ),
        32: "Education, Employment, and Civic Social Media Use as Moderators",
        33: (
            "Education and employment are longstanding predictors of civic participation because they shape civic skills, organizational exposure, and access to recruitment networks (Verba et al., 1995; Wilson, 2012). Higher education may expand awareness of volunteer opportunities and increase confidence in navigating organizational settings, while employment can provide routine social contact, institutional attachment, and recruitment opportunities. If in-person socialization supports volunteering partly by expanding access to networks, then education and employment may strengthen that association."
        ),
        34: (
            "Civic social media use is conceptually different. Digital tools may supplement volunteering by circulating invitations, information, and issue-based mobilization, but they may not fully replace face-to-face contact in organization-based participation. Prior work on connective action suggests that digitally mediated engagement can coexist with traditional civic action rather than simply substitute for it (Bennett & Segerberg, 2012). In this study, civic social media use is therefore treated as an additional moderator whose relationship to volunteering may vary across generations, especially if digital natives and older adults use online civic tools in different ways."
        ),
        37: "Data Source and Sample",
        38: (
            "Data were drawn from the Current Population Survey Civic Engagement and Volunteering Supplement (CPS-CEV), accessed through IPUMS CPS (Flood et al., 2023). The CPS-CEV is administered as a September supplement to the CPS in alternating years and is sponsored jointly by AmeriCorps and the U.S. Census Bureau. Because the CPS is based on a stratified multistage probability sample, the supplement provides nationally representative estimates of the U.S. civilian noninstitutionalized population."
        ),
        39: (
            "The pooled study sample includes respondents ages 18 and older from the 2017, 2019, 2021, and 2023 waves who had valid supplement weights (VLSUPPWT > 0) and nonmissing volunteering status (VLSTATUS in {1, 2}), yielding N = 201,168. Each wave contributes roughly 50,000 respondents. The pooled descriptive sample was used for the main analyses, while complete cases on the variables required for a given model were used within the regression, profile, and machine-learning stages. Because these data are repeated cross-sections rather than panel observations, generational comparisons should be interpreted as cohort-pattern differences rather than as direct evidence of individual change over time."
        ),
        40: "Measures",
        41: "Dependent Variable: Formal Volunteering",
        42: (
            "Formal volunteering was measured using VLSTATUS, which asks whether the respondent performed volunteer activities through or for an organization during the past 12 months. The variable was recoded as a binary indicator (1 = volunteered, 0 = did not volunteer). Focusing on formal volunteering keeps the dependent variable aligned with the organizational recruitment logic developed in the literature review."
        ),
        43: "Focal Independent Variable: In-Person Socialization",
        44: (
            "In-person socialization was measured with CESOCIALIZE, which asks how often respondents got together socially with friends, relatives, or neighbors during the past 12 months. Response options range from 1 (not at all) to 6 (basically every day). For the regression models, the variable was entered as a categorical factor with 'not at all' as the reference category so that potential threshold effects could be estimated without imposing linearity."
        ),
        45: "Generational Cohorts",
        46: (
            "Respondents were classified into five cohorts based on calculated birth year (survey year minus age): Generation Z (1997 or later), Millennials (1981-1996), Generation X (1965-1980), Baby Boomers (1946-1964), and the Silent Generation (before 1946). Because the CPS-CEV is fielded to adults, the age composition of Generation Z changes across waves; accordingly, same-age descriptive comparisons are treated as a robustness aid rather than as a full age-period-cohort test."
        ),
        47: "Moderating Variables",
        48: (
            "Three moderators were examined. Education was coded as bachelor's degree or higher versus less than a bachelor's degree using PEEDUCA. Employment status was coded as employed versus not employed using PEMLR. Civic social media use was measured with VLSOCMEDIA; for the three-way interaction models, this variable was operationalized as regular civic social media use versus infrequent or no use in order to preserve interpretable cell sizes across generation-by-socialization combinations."
        ),
        49: "Control Variables",
        50: (
            "All multivariable models adjusted for age, sex, race/ethnicity, marital status, household income, metropolitan status, region, and survey wave. These controls were included to reduce the risk that the estimated socialization-volunteering relationship simply reflected compositional differences across cohorts or across survey years."
        ),
        51: "Analytic Strategy",
        52: (
            "The study employed a sequential multi-method design. First, weighted descriptive statistics were used to document the distribution of in-person socialization across generations and survey waves and to establish whether the isolation patterns motivating the study were concentrated among particular cohorts."
        ),
        53: (
            "Second, survey-weighted logistic regression models were estimated to test the main hypotheses. Model 1 estimated the interaction between in-person socialization and generation. Models 2 through 4 estimated separate three-way interactions for education, employment, and civic social media use. Predicted probabilities and average marginal effects were used to interpret the magnitude of the socialization-volunteering relationship across cohorts. Same-age descriptive comparisons between Millennials and Generation Z were used only as a supplemental check on whether the observed pattern might reflect life-stage composition."
        ),
        54: (
            "Third, a person-centered latent profile analysis was used to identify broader patterns of civic engagement using five indicators: boycotting, contacting public officials, political conversation, in-person socialization, and organizational membership. The purpose of this stage was not to replace the regression analysis, but to assess whether the socialization-volunteering pattern emerged within a broader multidimensional structure of civic behavior and whether Generation Z was disproportionately concentrated in profiles marked by low social connection."
        ),
        55: (
            "Fourth, gradient boosting models with TreeSHAP were used to examine whether the regression findings were consistent with a nonlinear predictive model. This stage was designed to evaluate variable salience, inspect potential threshold patterns in in-person socialization, and compare whether the predictive importance of civic social media use and other correlates differed for Generation Z and older cohorts. TreeSHAP results were interpreted as model-based explanations of prediction rather than as causal effect estimates."
        ),
        56: (
            "All analyses were conducted in R 4.4 and Python 3.11. The regression models used supplement weights and linearized standard errors. The profile and machine-learning analyses were used to extend interpretation of the regression results, but the study's inferential claims remain associational given the repeated cross-sectional design."
        ),
    }

    for idx, text in revisions.items():
        set_paragraph_text(paragraphs[idx], text, highlight=True)
        clear_numbering(paragraphs[idx])

    method_paragraph = paragraphs[36]
    clear_numbering(method_paragraph)

    purpose_heading = add_paragraph_before(method_paragraph, "Purpose and Research Questions", style="Heading 2", highlight=True)
    purpose_text = add_paragraph_before(
        method_paragraph,
        "The purpose of this study is to examine whether in-person socialization is associated with formal volunteering, whether that association differs across generations, and whether education, employment, and civic social media use condition the relationship.",
        highlight=True,
    )
    rq1 = add_paragraph_before(
        method_paragraph,
        "Research Question 1: How does the relationship between in-person socialization frequency and formal volunteering differ across generational cohorts, and does this relationship exhibit nonlinear threshold effects?",
        highlight=True,
    )
    rq2 = add_paragraph_before(
        method_paragraph,
        "Research Question 2: To what extent do education, employment, and civic social media use moderate the association between in-person socialization and formal volunteering, and do these moderating effects differ by generation?",
        highlight=True,
    )
    hypotheses_heading = add_paragraph_before(method_paragraph, "Hypotheses", style="Heading 2", highlight=True)
    h1 = add_paragraph_before(
        method_paragraph,
        "Hypothesis 1: More frequent in-person socialization will be associated with a higher likelihood of formal volunteering.",
        highlight=True,
    )
    h2 = add_paragraph_before(
        method_paragraph,
        "Hypothesis 2: Generation Z will exhibit lower probabilities of formal volunteering than older cohorts across comparable levels of in-person socialization.",
        highlight=True,
    )
    h3 = add_paragraph_before(
        method_paragraph,
        "Hypothesis 3: Education and employment will strengthen the positive association between in-person socialization and formal volunteering.",
        highlight=True,
    )
    h4 = add_paragraph_before(
        method_paragraph,
        "Hypothesis 4: Civic social media use will moderate the association between in-person socialization and formal volunteering, but the size of that moderation will vary across generations.",
        highlight=True,
    )
    conceptual_heading = add_paragraph_before(method_paragraph, "Conceptual Model", style="Heading 2", highlight=True)
    conceptual_text = add_paragraph_before(
        method_paragraph,
        "Table 1 summarizes the study design. Formal volunteering is the dependent variable. In-person socialization is the focal independent variable, generation is modeled as a contextual moderator, education, employment, and civic social media use are modeled as additional moderators, and demographic and contextual characteristics are included as controls.",
        highlight=True,
    )
    caption = method_paragraph.insert_paragraph_before()
    clear_numbering(caption)
    caption.alignment = 1
    caption_run = caption.add_run(
        "Table 1. Conceptual model linking in-person socialization to formal volunteering across generations."
    )
    caption_run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    concept_table = insert_table_before(method_paragraph, rows=5, cols=3)
    concept_table.style = "Table Grid"
    cells = concept_table.rows[0].cells
    cells[0].text = "Construct"
    cells[1].text = "Role in Model"
    cells[2].text = "Operationalization"
    concept_rows = [
        ("Formal volunteering", "Dependent variable", "VLSTATUS recoded to volunteered vs. did not volunteer"),
        ("In-person socialization", "Focal independent variable", "CESOCIALIZE entered as six-category frequency measure"),
        ("Generation", "Contextual moderator", "Gen Z, Millennial, Gen X, Boomer, Silent cohorts"),
        ("Education, employment, civic social media use", "Additional moderators", "Resources and access conditions expected to alter the socialization-volunteering link"),
    ]
    for row_idx, values in enumerate(concept_rows, start=1):
        for col_idx, value in enumerate(values):
            concept_table.rows[row_idx].cells[col_idx].text = value

    doc.save(str(output_path))


if __name__ == "__main__":
    import sys

    out = Path(sys.argv[1]) if len(sys.argv) > 1 else Path("/tmp/revised_manuscript.docx")
    out.parent.mkdir(parents=True, exist_ok=True)
    revise_document(out)
