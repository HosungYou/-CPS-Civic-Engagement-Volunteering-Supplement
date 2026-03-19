#!/usr/bin/env python3
"""
Apply ALL revisions promised in Hosung's replies to Dr. Windon's comments.
Preserves all comment anchors. Yellow highlight = changed/new text.
"""
import zipfile, shutil, os, re, copy
from lxml import etree

SRC = os.path.expanduser("~/Downloads/Bowling_Alone_REVISED_with_comments.docx")
DST = os.path.expanduser("~/Downloads/Bowling_Alone_FINAL_REVISED.docx")
shutil.copy2(SRC, DST)

W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

z_in = zipfile.ZipFile(SRC, 'r')
doc_xml = z_in.read('word/document.xml')
tree = etree.fromstring(doc_xml)
body = tree.find(f'{{{W}}}body')


# ── Helpers ──────────────────────────────────────────────────────────
def get_text(elem):
    return ''.join(t.text or '' for t in elem.iter(f'{{{W}}}t'))


def find_para(fragment):
    for p in body.iter(f'{{{W}}}p'):
        if fragment in get_text(p):
            return p
    return None


def find_para_exact(text):
    for p in body.iter(f'{{{W}}}p'):
        if get_text(p).strip() == text:
            return p
    return None


def body_index(elem):
    for i, child in enumerate(body):
        if child is elem:
            return i
    return -1


def make_run(text, bold=False, italic=False, highlight=False, strike=False, fontsize=24):
    r = etree.Element(f'{{{W}}}r')
    rPr = etree.SubElement(r, f'{{{W}}}rPr')
    rF = etree.SubElement(rPr, f'{{{W}}}rFonts')
    rF.set(f'{{{W}}}ascii', 'Times New Roman')
    rF.set(f'{{{W}}}hAnsi', 'Times New Roman')
    sz = etree.SubElement(rPr, f'{{{W}}}sz')
    sz.set(f'{{{W}}}val', str(fontsize))
    szCs = etree.SubElement(rPr, f'{{{W}}}szCs')
    szCs.set(f'{{{W}}}val', str(fontsize))
    if bold:
        etree.SubElement(rPr, f'{{{W}}}b')
    if italic:
        etree.SubElement(rPr, f'{{{W}}}i')
    if highlight:
        hl = etree.SubElement(rPr, f'{{{W}}}highlight')
        hl.set(f'{{{W}}}val', 'yellow')
    if strike:
        etree.SubElement(rPr, f'{{{W}}}strike')
    t = etree.SubElement(r, f'{{{W}}}t')
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t.text = text
    return r


def make_para(text, bold=False, italic=False, highlight=False, center=False,
              indent=True, strike=False, fontsize=24):
    p = etree.Element(f'{{{W}}}p')
    pPr = etree.SubElement(p, f'{{{W}}}pPr')
    sp = etree.SubElement(pPr, f'{{{W}}}spacing')
    sp.set(f'{{{W}}}line', '480')
    sp.set(f'{{{W}}}lineRule', 'auto')
    if center:
        jc = etree.SubElement(pPr, f'{{{W}}}jc')
        jc.set(f'{{{W}}}val', 'center')
    if indent:
        ind = etree.SubElement(pPr, f'{{{W}}}ind')
        ind.set(f'{{{W}}}firstLine', '720')
    r = make_run(text, bold=bold, italic=italic, highlight=highlight,
                 strike=strike, fontsize=fontsize)
    p.append(r)
    return p


def make_multi_run_para(runs_data, center=False, indent=True):
    """runs_data: [(text, bold, italic, highlight), ...]"""
    p = etree.Element(f'{{{W}}}p')
    pPr = etree.SubElement(p, f'{{{W}}}pPr')
    sp = etree.SubElement(pPr, f'{{{W}}}spacing')
    sp.set(f'{{{W}}}line', '480')
    sp.set(f'{{{W}}}lineRule', 'auto')
    if center:
        jc = etree.SubElement(pPr, f'{{{W}}}jc')
        jc.set(f'{{{W}}}val', 'center')
    if indent:
        ind = etree.SubElement(pPr, f'{{{W}}}ind')
        ind.set(f'{{{W}}}firstLine', '720')
    for item in runs_data:
        text, bld, ita, hl = item[0], item[1], item[2], item[3]
        r = make_run(text, bold=bld, italic=ita, highlight=hl)
        p.append(r)
    return p


def replace_in_para(para, old, new, highlight=True):
    """Replace text across runs in a paragraph. Handles cross-run text."""
    # Collect all t elements with their parent runs
    t_elems = list(para.iter(f'{{{W}}}t'))
    full_text = ''.join(t.text or '' for t in t_elems)

    if old not in full_text:
        return False

    # Build a mapping: char_index -> (t_elem, offset_within_t)
    char_map = []
    for t in t_elems:
        txt = t.text or ''
        for j, ch in enumerate(txt):
            char_map.append((t, j))

    start = full_text.index(old)
    end = start + len(old)

    if start >= len(char_map) or end > len(char_map):
        return False

    # Find affected t elements
    first_t, first_offset = char_map[start]
    last_t, last_offset = char_map[end - 1]

    if first_t is last_t:
        # Simple case: replacement within a single text element
        orig = first_t.text or ''
        first_t.text = orig[:first_offset] + new + orig[last_offset + 1:]
        if highlight:
            parent_run = first_t.getparent()
            rPr = parent_run.find(f'{{{W}}}rPr')
            if rPr is None:
                rPr = etree.SubElement(parent_run, f'{{{W}}}rPr')
                parent_run.insert(0, rPr)
            existing_hl = rPr.find(f'{{{W}}}highlight')
            if existing_hl is None:
                hl_elem = etree.SubElement(rPr, f'{{{W}}}highlight')
                hl_elem.set(f'{{{W}}}val', 'yellow')
    else:
        # Cross-run: put replacement in first t, clear middle/end
        orig_first = first_t.text or ''
        first_t.text = orig_first[:first_offset] + new

        # Highlight the first run
        if highlight:
            parent_run = first_t.getparent()
            rPr = parent_run.find(f'{{{W}}}rPr')
            if rPr is None:
                rPr = etree.SubElement(parent_run, f'{{{W}}}rPr')
                parent_run.insert(0, rPr)
            existing_hl = rPr.find(f'{{{W}}}highlight')
            if existing_hl is None:
                hl_elem = etree.SubElement(rPr, f'{{{W}}}highlight')
                hl_elem.set(f'{{{W}}}val', 'yellow')

        # Clear text in subsequent affected t elements
        affected_ts = set()
        for i in range(start + 1, end):
            if i < len(char_map):
                affected_ts.add(char_map[i][0])
        affected_ts.discard(first_t)

        for t in affected_ts:
            if t is last_t:
                orig_last = t.text or ''
                t.text = orig_last[last_offset + 1:]
            else:
                t.text = ''

    return True


def apply_strikethrough(para):
    """Apply strikethrough to all runs in a paragraph."""
    for r in para.iter(f'{{{W}}}r'):
        rPr = r.find(f'{{{W}}}rPr')
        if rPr is None:
            rPr = etree.SubElement(r, f'{{{W}}}rPr')
            r.insert(0, rPr)
        existing = rPr.find(f'{{{W}}}strike')
        if existing is None:
            etree.SubElement(rPr, f'{{{W}}}strike')


def insert_after(ref_elem, new_elem):
    idx = body_index(ref_elem)
    if idx >= 0:
        body.insert(idx + 1, new_elem)
        return True
    return False


def insert_before(ref_elem, new_elem):
    idx = body_index(ref_elem)
    if idx >= 0:
        body.insert(idx, new_elem)
        return True
    return False


changes = []

# ════════════════════════════════════════════════════════════════════
# REPLY 1 → ABSTRACT TEXT CORRECTIONS (Comment 0)
# ════════════════════════════════════════════════════════════════════
abstract_p = find_para('Social isolation has been declared a public health epidemic')
if abstract_p is not None:
    r1 = replace_in_para(abstract_p,
        'yet its consequences for civic participation remain underexplored',
        'yet its implications for civic participation have received limited empirical attention')
    r2 = replace_in_para(abstract_p,
        'this study examines the relationship',
        'this study explores the association')
    r3 = replace_in_para(abstract_p,
        'Descriptive analyses reveal',
        'Descriptive analyses indicate')
    r4 = replace_in_para(abstract_p,
        'The transition from minimal socialization to even occasional social contact corresponds to the largest marginal gain in volunteering',
        'The transition from no socialization to even occasional social contact is associated with the largest observed difference in volunteering probability')
    r5 = replace_in_para(abstract_p,
        'Generation Z exhibits a volunteering ceiling',
        "Generation Z\u2019s predicted volunteering probability appears to level off")
    r6 = replace_in_para(abstract_p,
        "while education and civic social media use moderate isolation\u2019s civic consequences more effectively for older cohorts",
        'whereas education and civic social media use appear to moderate the socialization\u2013volunteering association more strongly among older cohorts')
    changes.append(f"Abstract: {sum([r1,r2,r3,r4,r5,r6])} text replacements")


# ════════════════════════════════════════════════════════════════════
# REPLY 3 → INTRODUCTION RESTRUCTURE (Comment 2)
# Problem → Gap → Purpose → Contribution
# ════════════════════════════════════════════════════════════════════

# 1. Tone down paragraph [18] — "striking generational asymmetry" → softer language
p18 = find_para('This omission is consequential')
if p18 is not None:
    replace_in_para(p18,
        'new evidence suggests that the erosion of social connection has entered a qualitatively different phase',
        'emerging evidence suggests that the erosion of social connection may have entered a qualitatively different phase')
    replace_in_para(p18,
        'reveal a striking generational asymmetry',
        'indicate a generational asymmetry in socialization patterns')
    changes.append("Intro [18]: toned down language")

# 2. Add GAP paragraph after para [19] (volunteering sector urgency)
p19 = find_para('Understanding the civic consequences of social isolation is particularly urgent')
if p19 is not None:
    gap_para = make_para(
        'Despite growing attention to the health consequences of social isolation '
        '(Cacioppo & Cacioppo, 2014; Holt-Lunstad et al., 2010), its implications for civic '
        'participation\u2014and particularly volunteering\u2014have received limited empirical '
        'scrutiny. Existing studies of generational differences in civic engagement have focused '
        'primarily on values and attitudes (Dalton, 2008; Zukin et al., 2006) rather than the '
        'social-structural conditions that enable or constrain participation. Moreover, the few '
        'studies linking social connectedness to volunteering have not examined whether this '
        'relationship operates differently across generational cohorts exposed to fundamentally '
        'different socialization environments.',
        highlight=True
    )
    insert_after(p19, gap_para)
    changes.append("Intro: inserted Gap paragraph after [19]")

# 3. Add PURPOSE paragraph after Gap
purpose_para = make_para(
    'The present study addresses this gap by using four waves of nationally representative '
    'CPS-CEV data (2017, 2019, 2021, 2023; N = 201,168) to examine the association between '
    'in-person socialization frequency and volunteering across generational cohorts, and to '
    'test whether education, employment, and civic social media use moderate this association '
    'differently by generation.',
    highlight=True
)
insert_after(gap_para, purpose_para)
changes.append("Intro: inserted Purpose paragraph")

# 4. Strikethrough old "The present study uses four waves..." paragraph [20]
p20 = find_para('The present study uses four waves of the CPS-CEV')
if p20 is not None:
    apply_strikethrough(p20)
    # Add note
    note_run = make_run(' [MOVED to Purpose and Research Questions section]',
                        bold=True, italic=True, highlight=True, fontsize=20)
    p20.append(note_run)
    changes.append("Intro [20]: strikethrough + moved note")

# 5. Strikethrough old RQ1 paragraph [21]
p21 = find_para('Research Question 1: How does the relationship between in-person socialization')
# Be careful: there are two RQ1 paragraphs (original + new section). Target the one in intro.
for p in body.iter(f'{{{W}}}p'):
    txt = get_text(p)
    if 'Research Question 1: How does the relationship between in-person socialization' in txt:
        # Check if it's in the intro (not the new Purpose section) by checking for highlight
        hl_elems = list(p.iter(f'{{{W}}}highlight'))
        if not hl_elems:  # Original (no highlight) = intro version
            apply_strikethrough(p)
            note = make_run(' [MOVED to Purpose and Research Questions section]',
                            bold=True, italic=True, highlight=True, fontsize=20)
            p.append(note)
            changes.append("Intro RQ1: strikethrough + moved note")
            break

# 6. Strikethrough old RQ2 paragraph [22]
for p in body.iter(f'{{{W}}}p'):
    txt = get_text(p)
    if 'Research Question 2: To what extent do education, employment, and civic social media' in txt:
        hl_elems = list(p.iter(f'{{{W}}}highlight'))
        if not hl_elems:
            apply_strikethrough(p)
            note = make_run(' [MOVED to Purpose and Research Questions section]',
                            bold=True, italic=True, highlight=True, fontsize=20)
            p.append(note)
            changes.append("Intro RQ2: strikethrough + moved note")
            break

# 7. Tone down contribution paragraph [23]
p23 = find_para('By addressing these questions with nationally representative data')
if p23 is not None:
    replace_in_para(p23,
        'it provides the first large-scale empirical evidence connecting',
        'it extends the Surgeon General\u2019s health-focused framing to civic outcomes, examining the association between')
    replace_in_para(p23,
        'identifies the transition from social isolation to even minimal social contact as the most consequential threshold for volunteering participation',
        'identifies the transition from social isolation to even minimal social contact as a potentially consequential threshold for volunteering participation')
    replace_in_para(p23,
        'it demonstrates that conventional pathways expected to promote civic participation',
        'it examines whether conventional pathways expected to promote civic participation')
    replace_in_para(p23,
        'with substantially weaker moderating effects for Generation Z',
        'may operate differently across generations, with potentially weaker associations for Generation Z')
    changes.append("Intro [23]: contribution paragraph toned down")


# ════════════════════════════════════════════════════════════════════
# REPLY 5 → RENAME "Theoretical Framework" → "Literature Review" (Comment 4)
# ════════════════════════════════════════════════════════════════════
tf_heading = find_para_exact('Theoretical Framework')
if tf_heading is not None:
    for t in tf_heading.iter(f'{{{W}}}t'):
        if t.text and 'Theoretical Framework' in t.text:
            t.text = t.text.replace('Theoretical Framework', 'Literature Review')
    # Add highlight
    for r in tf_heading.iter(f'{{{W}}}r'):
        rPr = r.find(f'{{{W}}}rPr')
        if rPr is None:
            rPr = etree.SubElement(r, f'{{{W}}}rPr')
            r.insert(0, rPr)
        hl = etree.SubElement(rPr, f'{{{W}}}highlight')
        hl.set(f'{{{W}}}val', 'yellow')
    changes.append("Heading: 'Theoretical Framework' → 'Literature Review'")


# ════════════════════════════════════════════════════════════════════
# REPLY 7 → LIT REVIEW: ADD NEW SUBSECTIONS (Comment 6)
# 5 focused subsections matching RQ keywords
# ════════════════════════════════════════════════════════════════════

# (A) After existing "Generational Shifts" content [31], add "Why Gen Z" content
p31 = find_para('These theoretical perspectives generate competing predictions')
if p31 is not None:
    # Insert "Volunteering Across Generations: Why Generation Z?" heading
    why_genz_heading = make_para(
        'Volunteering Across Generations: Why Generation Z?',
        bold=True, highlight=True, indent=False, center=False
    )
    insert_after(p31, why_genz_heading)

    why_genz_1 = make_para(
        'Generation Z is the focal cohort in this study for three reasons. First, Gen Z is '
        'the first generation to have entered adulthood entirely within a digitally mediated '
        'social environment, making them a critical test case for whether digital connectivity '
        'substitutes for or undermines face-to-face civic recruitment (Bennett & Segerberg, 2012). '
        'Second, members of Gen Z entered adulthood during or immediately before the COVID-19 '
        'pandemic, compounding existing socialization deficits during a formative period for '
        'civic habit formation (Flanagan & Levine, 2010). Third, exploratory analyses of the '
        'CPS-CEV data indicate that Gen Z exhibits the highest rates of social isolation among '
        'all generational cohorts\u2014nearly half report rarely or never socializing in person\u2014'
        'raising the question of whether the well-established socialization\u2013volunteering '
        'association operates differently for this generation.',
        highlight=True
    )
    insert_after(why_genz_heading, why_genz_1)

    why_genz_2 = make_para(
        'Generation is operationalized as a moderating variable rather than merely a demographic '
        'control because generational membership captures shared formative experiences\u2014including '
        'exposure to digital technology, economic conditions, and political events during '
        'adolescence\u2014that may fundamentally alter how social interaction translates into civic '
        'behavior (Hooghe, 2012). Prior research supports this expectation: Kelle et al. (2025) '
        'demonstrated that Baby Boomers maintained higher voluntary engagement than subsequent '
        'cohorts even after accounting for age effects, suggesting that generational context shapes '
        'not only the level but also the form of civic participation.',
        highlight=True
    )
    insert_after(why_genz_1, why_genz_2)
    changes.append("Lit Review: inserted 'Why Gen Z?' subsection with 2 paragraphs")

    # (B) Add "Education, Employment, and Volunteering" subsection
    edu_heading = make_para(
        'Education, Employment, and Volunteering',
        bold=True, highlight=True, indent=False, center=False
    )
    insert_after(why_genz_2, edu_heading)

    edu_para1 = make_para(
        'The civic voluntarism model identifies education and employment as primary resource '
        'predictors of civic participation (Verba et al., 1995). Higher education develops '
        'civic skills, expands social networks, and increases exposure to recruitment appeals '
        '(Brady et al., 1995; Schlozman et al., 2012). Employment provides organizational '
        'membership, collegial networks, and economic resources that facilitate volunteering '
        '(Musick & Wilson, 2008; Wilson, 2000). These pathways are well-documented across '
        'decades of research, with education consistently emerging as the strongest predictor '
        'of volunteering (Wilson, 2012).',
        highlight=True
    )
    insert_after(edu_heading, edu_para1)

    edu_para2 = make_para(
        'However, whether these pathways operate with equal strength across generational cohorts '
        'has received limited attention. If Generation Z\u2019s social isolation reflects a structural '
        'shift in how young adults form social connections\u2014rather than a deficit in educational '
        'or employment resources\u2014then the conventional expectation that higher education and '
        'employment will promote volunteering may not hold equally across generations. This study '
        'examines education and employment as moderators of the socialization\u2013volunteering '
        'association to test this possibility.',
        highlight=True
    )
    insert_after(edu_para1, edu_para2)
    changes.append("Lit Review: inserted 'Education, Employment, and Volunteering' subsection")

    # (C) Add "Civic Social Media Use as a Moderator" subsection
    sm_heading = make_para(
        'Civic Social Media Use as a Moderator',
        bold=True, highlight=True, indent=False, center=False
    )
    insert_after(edu_para2, sm_heading)

    sm_para = make_para(
        'We additionally examine civic use of social media as a potential moderating factor. '
        'Bennett and Segerberg (2012) theorized the emergence of \u201cconnective action\u201d\u2014'
        'individualized, digitally mediated civic expression\u2014as a potential complement or '
        'substitute for traditional collective action. If digital tools have genuinely created '
        'alternative pathways to civic recruitment, then civic social media use should compensate '
        'for the absence of face-to-face interaction, particularly for digital natives. However, '
        'if social media functions as ambient noise for a generation that has never known life '
        'without it, the compensatory effect may be weaker for Generation Z than for older adults '
        'who adopted digital tools as a novel civic channel.',
        highlight=True
    )
    insert_after(sm_heading, sm_para)
    changes.append("Lit Review: inserted 'Civic Social Media' subsection")


# ════════════════════════════════════════════════════════════════════
# REPLY 11 → RENAME "Democratic Health" → "Civic Consequences" (Comment 10)
# ════════════════════════════════════════════════════════════════════
dem_heading = find_para('The Loneliness Epidemic and Democratic Health')
if dem_heading is not None:
    for t in dem_heading.iter(f'{{{W}}}t'):
        if t.text and 'Democratic Health' in t.text:
            t.text = t.text.replace('Democratic Health', 'Civic Consequences')
    for r in dem_heading.iter(f'{{{W}}}r'):
        rPr = r.find(f'{{{W}}}rPr')
        if rPr is None:
            rPr = etree.SubElement(r, f'{{{W}}}rPr')
            r.insert(0, rPr)
        hl = etree.SubElement(rPr, f'{{{W}}}highlight')
        hl.set(f'{{{W}}}val', 'yellow')
    changes.append("Heading: 'Democratic Health' → 'Civic Consequences'")

# Also replace "democratic health" in body text if it appears
for p in body.iter(f'{{{W}}}p'):
    txt = get_text(p)
    if 'democratic health' in txt.lower() and 'Democratic Health' not in txt:
        replace_in_para(p, 'democratic health', 'civic participation')


# ════════════════════════════════════════════════════════════════════
# REPLY 9 → MOVE LONELINESS FRAMING TO INTRO (Comment 8)
# Add a bridging sentence at the end of the loneliness section
# pointing to the intro, and add loneliness content to intro
# ════════════════════════════════════════════════════════════════════
# Modify the intro's first paragraph to integrate the loneliness framing more
p17 = find_para('In May 2023, the U.S. Surgeon General declared')
if p17 is not None:
    replace_in_para(p17,
        'Yet the advisory focused almost exclusively on health outcomes. The civic consequences of the loneliness epidemic',
        'Yet the advisory\u2019s empirical evidence focused overwhelmingly on health outcomes, leaving the civic consequences of social disconnection')
    changes.append("Intro [17]: integrated loneliness framing language")


# ════════════════════════════════════════════════════════════════════
# REPLY 17 → METHOD SECTION (Comment 16)
# Add "Why CPS-CEV fits" justification
# ════════════════════════════════════════════════════════════════════
data_source_para = find_para('Data come from the Current Population Survey')
if data_source_para is not None:
    # Add justification sentence at the end of data source paragraph
    fit_para = make_para(
        'The CPS-CEV is uniquely suited to this study because it is the only nationally '
        'representative survey that simultaneously measures both volunteering behavior and '
        'in-person socialization frequency across multiple waves, enabling examination of '
        'the socialization\u2013volunteering association over time and across the COVID-19 pandemic.',
        highlight=True
    )
    insert_after(data_source_para, fit_para)
    changes.append("Method: added 'Why CPS-CEV fits' justification")

# Add missing data handling paragraph after sample paragraph
sample_para = find_para('We pooled four waves of CPS-CEV data')
if sample_para is not None:
    missing_para = make_para(
        'Respondents with missing values on the key independent variable (CESOCIALIZE) were '
        'excluded from models involving socialization (less than 1% of the analytic sample). '
        'For moderating variables, cases with missing values were excluded from the respective '
        'interaction models. All supplement weights (VLSUPPWT) were applied to produce nationally '
        'representative estimates.',
        highlight=True
    )
    insert_after(sample_para, missing_para)
    changes.append("Method: added missing data handling paragraph")


# ════════════════════════════════════════════════════════════════════
# COMMENT 18 → Shorten coding explanation, sub-sections, reduce jargon
# Strikethrough old supplementary analysis paragraph [78]
# (since new justification already exists at [75-77])
# ════════════════════════════════════════════════════════════════════
old_suppl = find_para('As a supplementary analysis, we employ Latent Profile Analysis (LPA)')
if old_suppl is not None:
    # Check it's the OLD one (has comment anchor for 18)
    has_comment = old_suppl.findall(f'.//{{{W}}}commentRangeStart')
    if has_comment:
        apply_strikethrough(old_suppl)
        note = make_run(
            ' [REPLACED: See revised Supplementary Analyses section above with simplified terminology]',
            bold=True, italic=True, highlight=True, fontsize=20
        )
        old_suppl.append(note)
        changes.append("Method: strikethrough old supplementary paragraph (replaced by new justification)")

# Shorten the coding note in the Key Independent Variable paragraph
coding_para = find_para('In-person socialization frequency was measured using CESOCIALIZE')
if coding_para is not None:
    # The long coding note starts with "An important coding note:"
    replace_in_para(coding_para,
        'We note that the \u201cnot at all\u201d response category captures respondents who report no in-person social gatherings over a 12-month period. While some respondents selecting this option may have had incidental social contact, the response represents the lowest point on the socialization continuum and is interpreted as indicating minimal or absent intentional social engagement. An important coding note: the civic engagement supplement variables (CE*) use ascending frequency (1 = not at all to 6 = basically every day), whereas volunteering supplement variables (VL*) use descending frequency. All coding was verified through education cross-validation, confirming that higher socialization frequencies correspond to higher educational attainment, as expected from the civic voluntarism model.',
        'The \u201cnot at all\u201d category represents the lowest point on the socialization continuum and is interpreted as indicating minimal or absent intentional social engagement.',
        highlight=True)
    # Add footnote-style note as separate paragraph
    coding_note = make_para(
        'Note: The civic engagement supplement variables (CE*) use ascending frequency coding '
        '(1 = not at all to 6 = basically every day), whereas the volunteering supplement '
        'variables (VL*) use descending frequency. All coding was verified through education '
        'cross-validation.',
        highlight=True, italic=True
    )
    insert_after(coding_para, coding_note)
    changes.append("Method: shortened coding explanation, moved detail to note")


# ════════════════════════════════════════════════════════════════════
# SAVE
# ════════════════════════════════════════════════════════════════════
modified_xml = etree.tostring(tree, xml_declaration=True, encoding='UTF-8', standalone=True)

with zipfile.ZipFile(DST, 'w', zipfile.ZIP_DEFLATED) as z_out:
    for item in z_in.namelist():
        if item == 'word/document.xml':
            z_out.writestr(item, modified_xml)
        else:
            z_out.writestr(item, z_in.read(item))

z_in.close()

# Re-embed the figure (it was in the source)
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

FIG = os.path.expanduser("~/Downloads/figure1_conceptual_model.png")
if os.path.exists(FIG):
    ddoc = Document(DST)
    for p in ddoc.paragraphs:
        if '[See Figure 1 inserted below]' in p.text:
            for run in p.runs:
                run.text = ''
            run = p.add_run()
            run.add_picture(FIG, width=Inches(6.0))
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            print("Figure 1 re-embedded")
            break
    ddoc.save(DST)

print(f"\n{'='*60}")
print(f"All revisions applied. Changes made:")
for c in changes:
    print(f"  - {c}")
print(f"\nSaved: {DST}")
print(f"Size: {os.path.getsize(DST):,} bytes")
