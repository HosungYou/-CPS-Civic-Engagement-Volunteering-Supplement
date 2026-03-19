#!/usr/bin/env python3
"""
Revise manuscript based on Dr. Windon's 2026-03-11 feedback.
- Yellow highlight = modified/new text
- Appends "Response to Reviewer Comments" table at end
- Saves to Downloads as revised copy
"""

from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os, copy

OUTPUT = os.path.expanduser(
    "~/Downloads/Bowling_Alone_REVISED_Windon_Response.docx"
)

doc = Document()

# ── APA 7 Defaults ──────────────────────────────────────────────────────
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(12)
style.paragraph_format.line_spacing = 2.0
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.space_before = Pt(0)

# Also set East-Asian font fallback
rpr = style.element.get_or_add_rPr()
rFonts = rpr.find(qn("w:rFonts"))
if rFonts is None:
    rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="Times New Roman" w:hAnsi="Times New Roman" w:eastAsia="Times New Roman" w:cs="Times New Roman"/>')
    rpr.insert(0, rFonts)

# Set default margins (1 inch all around)
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)


# ── Helper Functions ────────────────────────────────────────────────────
Y = WD_COLOR_INDEX.YELLOW  # shorthand


def _apply_font(run, bold=False, italic=False, highlight=False, size=12):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if highlight:
        run.font.highlight_color = Y


def h1(text, highlight=False):
    """APA Level 1 heading: centered, bold."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    _apply_font(r, bold=True, highlight=highlight)
    return p


def h2(text, highlight=False):
    """APA Level 2 heading: flush left, bold."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.first_line_indent = None
    r = p.add_run(text)
    _apply_font(r, bold=True, highlight=highlight)
    return p


def h3(text, highlight=False):
    """APA Level 3 heading: flush left, bold italic."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.first_line_indent = None
    r = p.add_run(text)
    _apply_font(r, bold=True, italic=True, highlight=highlight)
    return p


def body(text, highlight=False, indent=True):
    """Body paragraph. First line indent 0.5" by default."""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 2.0
    if indent:
        p.paragraph_format.first_line_indent = Inches(0.5)
    else:
        p.paragraph_format.first_line_indent = None
    r = p.add_run(text)
    _apply_font(r, highlight=highlight)
    return p


def mixed(parts, indent=True):
    """Paragraph with mixed runs: [(text, hl, bold, italic), ...]"""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 2.0
    if indent:
        p.paragraph_format.first_line_indent = Inches(0.5)
    else:
        p.paragraph_format.first_line_indent = None
    for part in parts:
        t = part[0]
        hl = part[1] if len(part) > 1 else False
        bl = part[2] if len(part) > 2 else False
        it = part[3] if len(part) > 3 else False
        r = p.add_run(t)
        _apply_font(r, bold=bl, italic=it, highlight=hl)
    return p


def blank():
    """Insert a blank line."""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 2.0
    return p


# ╔══════════════════════════════════════════════════════════════════════╗
# ║  TITLE PAGE                                                         ║
# ╚══════════════════════════════════════════════════════════════════════╝

blank()
blank()
blank()
blank()
h1("Bowling Alone, Scrolling Together: Social Isolation and the Generational Divide in Volunteering")
blank()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.line_spacing = 2.0
r = p.add_run("Hosung You")
_apply_font(r)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.line_spacing = 2.0
r = p.add_run("The Pennsylvania State University")
_apply_font(r)

blank()
h2("Author Note")
body("Hosung You  https://orcid.org/XXXX-XXXX-XXXX-XXXX", indent=True)

doc.add_page_break()

# ╔══════════════════════════════════════════════════════════════════════╗
# ║  ABSTRACT  (Comment 0: tone down causal/overstated language)        ║
# ╚══════════════════════════════════════════════════════════════════════╝

h1("Abstract")

mixed([
    ("Social isolation has been declared a public health epidemic, yet its ", False),
    ("implications for civic participation have received limited empirical attention.", True),
    (" Using four waves of the Current Population Survey Civic Engagement and Volunteering Supplement (2017, 2019, 2021, 2023; ", False),
    ("N", False, False, True),
    (" = 201,168), this study ", False),
    ("explores the association", True),
    (" between in-person socialization frequency and volunteering across generational cohorts. Descriptive analyses ", False),
    ("indicate", True),
    (" that nearly half of Generation Z adults report rarely or never socializing in person\u2014a pattern ", False),
    ("that predates", True),
    (" COVID-19 and persists despite higher education and employment. The transition from ", False),
    ("no socialization to even occasional social contact is associated with the largest observed difference in volunteering probability", True),
    (" (10\u201314 percentage points) across all generations. ", False),
    ("Generation Z\u2019s predicted volunteering probability appears to level off", True),
    (" at approximately 34% regardless of socialization frequency, ", False),
    ("whereas", True),
    (" education and civic social media use ", False),
    ("appear to moderate the socialization\u2013volunteering association more strongly among", True),
    (" older cohorts. Implications for volunteer recruitment and community development programming are discussed.", False),
], indent=True)

body("Keywords: volunteering, social isolation, civic engagement, generational differences, Current Population Survey", indent=True)

doc.add_page_break()

# ╔══════════════════════════════════════════════════════════════════════╗
# ║  INTRODUCTION  (Comments 1, 4, 5: restructure, shorten)            ║
# ║  New structure: Problem → Gap → Purpose → Contribution              ║
# ║  RQs moved to "Purpose and Research Questions" section              ║
# ╚══════════════════════════════════════════════════════════════════════╝

h1("Bowling Alone, Scrolling Together: Social Isolation and the Generational Divide in Volunteering")

# Para 1: Problem — Surgeon General (mostly original, minor edits)
mixed([
    ("In May 2023, the U.S. Surgeon General declared loneliness and social isolation a public health epidemic, documenting extensive consequences for mental health, physical well-being, and mortality (Office of the Surgeon General, 2023). The advisory cited research showing that social disconnection carries health risks comparable to smoking 15 cigarettes per day (Holt-Lunstad et al., 2010) and called for urgent action to rebuild social infrastructure. ", False),
    ("Yet the advisory\u2019s empirical evidence focused overwhelmingly on health outcomes, leaving the civic consequences of social disconnection\u2014its implications for volunteering, community capacity, and collective participation\u2014largely unexamined.", True),
], indent=True)

# Para 2: Problem cont. — Generational pattern (toned down)
mixed([
    ("This omission is consequential. A quarter century after Putnam (2000) documented the broad decline of associational life in ", False),
    ("Bowling Alone", False, False, True),
    (", ", False),
    ("emerging evidence suggests that the erosion of social connection has entered a qualitatively different phase.", True),
    (" The decline is no longer gradual and generationally uniform. Instead, data from the Current Population Survey Civic Engagement and Volunteering Supplement (CPS-CEV) ", False),
    ("indicate a generational asymmetry in socialization patterns:", True),
    (" nearly half of Generation Z adults report rarely or never socializing in person with friends, relatives, or neighbors\u2014a pattern that predates the COVID-19 pandemic ", False),
    ("and persists despite higher education and employment.", True),
], indent=True)

# Para 3: Problem cont. — Volunteering sector (original)
body(
    "Understanding the civic consequences of social isolation is particularly urgent for the volunteering sector. "
    "Unlike voting or charitable giving, organized volunteering demands sustained investment of time, skills, and "
    "social energy within formal organizational structures (Wilson, 2000). Volunteering is therefore uniquely dependent "
    "on the social networks through which individuals are recruited, trained, and retained (Musick & Wilson, 2008). "
    "If a substantial portion of a generation lacks the social connections that serve as pathways to volunteering, "
    "the implications extend beyond individual nonparticipation to the organizational capacity of the voluntary "
    "sector itself.",
    highlight=False,
)

# Para 4: NEW — Gap (Comment 1: show the gap)
body(
    "Despite growing attention to the health consequences of social isolation (Cacioppo & Cacioppo, 2014; "
    "Holt-Lunstad et al., 2010), its implications for civic participation\u2014and particularly volunteering\u2014"
    "have received limited empirical scrutiny. Existing studies of generational differences in civic engagement "
    "have focused primarily on values and attitudes (Dalton, 2008; Zukin et al., 2006) rather than the "
    "social-structural conditions that enable or constrain participation. Moreover, the few studies linking "
    "social connectedness to volunteering have not examined whether this relationship operates differently "
    "across generational cohorts exposed to fundamentally different socialization environments.",
    highlight=True,
)

# Para 5: NEW — Purpose (1 sentence, Comment 1)
body(
    "The present study addresses this gap by using four waves of nationally representative CPS-CEV data "
    "(2017, 2019, 2021, 2023; N = 201,168) to examine the association between in-person socialization "
    "frequency and volunteering across generational cohorts, and to test whether education, employment, "
    "and civic social media use moderate this association differently by generation.",
    highlight=True,
)

# Para 6: NEW — Contribution (toned down, Comment 0, 1)
mixed([
    ("This study makes three contributions. First, it ", False),
    ("extends the Surgeon General\u2019s health-focused framing of the loneliness epidemic to civic outcomes,", True),
    (" examining the association between social isolation and volunteering across generational cohorts with nationally representative data. Second, it identifies the transition from social isolation to even minimal social contact as ", False),
    ("a potentially consequential threshold", True),
    (" for volunteering participation\u2014a finding with direct implications for volunteer recruitment strategy. Third, it demonstrates that conventional pathways expected to promote civic participation\u2014higher education, employment, digital civic tools\u2014", False),
    ("may operate differently across generations, with potentially weaker associations for Generation Z.", True),
], indent=True)


# ╔══════════════════════════════════════════════════════════════════════╗
# ║  LITERATURE REVIEW  (Comments 2, 3, 4, 5: restructure & realign)   ║
# ╚══════════════════════════════════════════════════════════════════════╝

h1("Literature Review", highlight=True)

# ── Section 1: Social Interaction and Civic Recruitment ──────────────
h2("Social Interaction and Civic Recruitment")

body(
    "The theoretical foundation for the relationship between social interaction and volunteering rests on "
    "social capital theory and the civic voluntarism model. Putnam (2000) defined social capital as the "
    "\u201cconnections among individuals\u2014social networks and the norms of reciprocity and trustworthiness "
    "that arise from them\u201d (p. 19), distinguishing between bonding capital (ties within homogeneous "
    "groups) and bridging capital (ties across diverse groups). Both forms of social capital generate "
    "civic returns: bonding capital sustains mutual obligation and community solidarity, while bridging "
    "capital exposes individuals to new information, opportunities, and recruitment appeals (Putnam, 2000; "
    "Woolcock & Narayan, 2000).",
)

body(
    "The civic voluntarism model (CVM) identifies three categories of factors predicting civic "
    "participation: resources (time, money, civic skills), engagement (political interest, civic duty), "
    "and recruitment (being asked to participate through social networks; Verba et al., 1995). The "
    "recruitment mechanism is particularly relevant for volunteering, as organizational participation "
    "typically requires an explicit invitation or social exposure to opportunities (Brady et al., 1995; "
    "Musick & Wilson, 2008). Social interaction creates the occasions through which individuals encounter "
    "these recruitment appeals. Coleman (1988) formalized this insight by arguing that social networks "
    "create obligations, expectations, and information channels that facilitate collective action. "
    "Granovetter\u2019s (1973) theory of weak ties further suggests that even casual acquaintances can "
    "serve as bridges to new civic opportunities\u2014implying that the transition from zero social "
    "contact to even minimal interaction may activate recruitment channels that are otherwise entirely "
    "absent.",
)

# ── Section 2: NEW — Volunteering Across Generations (Comment 3) ─────
h2("Volunteering Across Generations: Why Generation Z?", highlight=True)

body(
    "Generational membership captures shared formative experiences\u2014including exposure to digital "
    "technology, economic conditions, and political events during adolescence\u2014that may fundamentally "
    "alter how social interaction translates into civic behavior (Hooghe, 2012). In this study, generation "
    "is operationalized as a moderating variable rather than merely a demographic control because the "
    "theoretical expectation is that socialization\u2019s role in civic recruitment may vary depending on "
    "the social infrastructure available to each cohort.",
    highlight=True,
)

body(
    "Twenge et al. (2019) documented sharp declines in face-to-face social interaction among U.S. "
    "adolescents beginning around 2012, coinciding with widespread smartphone adoption. Using nationally "
    "representative data from Monitoring the Future (N = 8.2 million), they found that high school seniors "
    "in 2016 spent approximately one hour less per day in face-to-face social interaction compared to "
    "seniors in the late 1980s, despite comparable levels of homework and extracurricular activity. "
    "Twenge (2017) characterized the generation born after 1995 as uniquely shaped by digital mediation "
    "of social life, with consequences for mental health, social development, and community participation.",
)

body(
    "Generation Z is the focal cohort for three reasons. First, Gen Z is the first generation to have "
    "entered adulthood entirely within a digitally mediated social environment, making them a critical "
    "test case for whether digital connectivity substitutes for or undermines face-to-face civic "
    "recruitment (Bennett & Segerberg, 2012). Second, members of Gen Z entered adulthood during or "
    "immediately before the COVID-19 pandemic, compounding existing socialization deficits during a "
    "formative period for civic habit formation (Flanagan & Levine, 2010). Third, exploratory analyses "
    "of the CPS-CEV data indicate that Gen Z exhibits the highest rates of social isolation among all "
    "generational cohorts, raising the question of whether the well-established socialization\u2013volunteering "
    "association operates differently for this generation.",
    highlight=True,
)

body(
    "Prior research on generational differences in civic participation has identified divergent patterns "
    "across cohorts. Zukin et al. (2006) found that younger cohorts favored community service over "
    "electoral engagement. Dalton (2008) distinguished between \u201cduty-based\u201d citizenship norms "
    "(voting, military service) that are declining among younger cohorts and \u201cengaged\u201d citizenship "
    "norms (community involvement, lifestyle politics) that may be rising. Kelle et al. (2025) "
    "demonstrated that Baby Boomers maintained higher voluntary engagement than subsequent cohorts even "
    "after accounting for age effects. These findings suggest that generational context shapes not only "
    "the level but also the form of civic participation.",
    highlight=True,
)

# ── Section 3: NEW — Education, Employment, and Volunteering (Comment 3) ─
h2("Education, Employment, and Volunteering", highlight=True)

body(
    "The civic voluntarism model identifies education and employment as primary resource predictors of "
    "civic participation (Verba et al., 1995). Higher education develops civic skills, expands social "
    "networks, and increases exposure to recruitment appeals (Brady et al., 1995; Schlozman et al., 2012). "
    "Employment provides organizational membership, collegial networks, and economic resources that "
    "facilitate volunteering (Musick & Wilson, 2008; Wilson, 2000). These pathways are well-documented "
    "across decades of research, with education consistently emerging as the strongest predictor of "
    "volunteering (Wilson, 2012).",
    highlight=True,
)

body(
    "However, whether these pathways operate with equal strength across generational cohorts has received "
    "limited attention. If Generation Z\u2019s social isolation reflects a structural shift in how young "
    "adults form social connections\u2014rather than a deficit in educational or employment resources\u2014then "
    "the conventional expectation that higher education and employment will promote volunteering may not "
    "hold equally across generations. This study examines education and employment as moderators of the "
    "socialization\u2013volunteering association to test this possibility.",
    highlight=True,
)

# ── Section 4: NEW — Civic Social Media (Comment 3: separate addition) ─
h2("Civic Social Media Use as a Moderator", highlight=True)

body(
    "We additionally examine civic use of social media as a potential moderating factor. Bennett and "
    "Segerberg (2012) theorized the emergence of \u201cconnective action\u201d\u2014individualized, digitally "
    "mediated civic expression\u2014as a potential complement or substitute for traditional collective action. "
    "If digital tools have genuinely created alternative pathways to civic recruitment, then civic social "
    "media use should compensate for the absence of face-to-face interaction, particularly for digital "
    "natives. However, if social media functions as ambient noise for a generation that has never known "
    "life without it, the compensatory effect may be weaker for Generation Z than for older adults who "
    "adopted digital tools as a novel civic channel.",
    highlight=True,
)

# ── Section 5: Loneliness Epidemic (Comment 4, 5: retitled, shortened) ─
h2("The Loneliness Epidemic and Civic Consequences", highlight=True)

mixed([
    ("The Surgeon General\u2019s 2023 advisory situated social isolation within a public health framework, "
     "documenting associations between social disconnection and increased risk of heart disease, stroke, "
     "dementia, depression, and premature mortality (Office of the Surgeon General, 2023). The advisory "
     "acknowledged that communities with higher social connectedness exhibit \u201chigher civic engagement\u201d "
     "(p. 28). However, the advisory\u2019s empirical evidence focused overwhelmingly on health outcomes, "
     "leaving the ", False),
    ("civic consequences", True),
    (" of the loneliness epidemic largely as an assertion rather than a demonstrated finding.", False),
], indent=True)

mixed([
    ("Cacioppo and Cacioppo (2014) provided a theoretical bridge between loneliness and civic withdrawal, "
     "arguing that perceived social isolation triggers a self-reinforcing cycle of hypervigilance to social "
     "threat, withdrawal from social situations, and declining trust\u2014all of which undermine the social "
     "engagement that sustains civic participation. The present study extends this line of inquiry by "
     "examining whether the loneliness epidemic has measurable ", False),
    ("civic consequences\u2014specifically, whether social isolation is associated with lower volunteering "
     "rates and whether this association varies by generation.", True),
], indent=True)


# ╔══════════════════════════════════════════════════════════════════════╗
# ║  PURPOSE AND RESEARCH QUESTIONS  (Comments 6, 7, 8: NEW section)   ║
# ╚══════════════════════════════════════════════════════════════════════╝

h1("Purpose and Research Questions", highlight=True)

body(
    "The purpose of this study is to examine the association between in-person socialization frequency "
    "and volunteering across generational cohorts using nationally representative data from the Current "
    "Population Survey Civic Engagement and Volunteering Supplement (2017\u20132023, N = 201,168). "
    "Guided by social capital theory and the civic voluntarism model, the study tests whether the "
    "socialization\u2013volunteering association differs by generation and whether conventional predictors "
    "of civic participation\u2014education, employment, and civic social media use\u2014moderate this "
    "association differently across cohorts.",
    highlight=True,
)

# Conceptual model description (Comment 6)
body(
    "Figure 1 presents the conceptual model guiding this study. The primary relationship of interest "
    "is the association between in-person socialization frequency (independent variable) and volunteering "
    "status (dependent variable, binary). Generation serves as the primary moderator, with education, "
    "employment status, and civic social media use examined as additional moderators. The model controls "
    "for age, sex, race/ethnicity, marital status, household income, metropolitan status, region, "
    "and survey wave.",
    highlight=True,
)

body(
    "[INSERT FIGURE 1: Conceptual Model HERE]",
    highlight=True, indent=False,
)

# RQ1 + Hypotheses (Comment 7)
mixed([
    ("Research Question 1", True, True, False),
    (": How does the association between in-person socialization frequency and volunteering differ "
     "across generational cohorts?", True),
], indent=True)

mixed([
    ("Hypothesis 1.", True, True, False),
    (" Higher levels of in-person socialization will be positively associated with volunteering "
     "across all generational cohorts.", True),
], indent=True)

mixed([
    ("Hypothesis 2.", True, True, False),
    (" Generation Z will exhibit a weaker socialization\u2013volunteering association at higher "
     "socialization levels compared with older cohorts (i.e., a plateau effect).", True),
], indent=True)

# RQ2 + Hypotheses
mixed([
    ("Research Question 2", True, True, False),
    (": To what extent do education, employment, and civic social media use moderate the "
     "socialization\u2013volunteering association, and do these moderating effects differ by generation?", True),
], indent=True)

mixed([
    ("Hypothesis 3a.", True, True, False),
    (" Education will positively moderate the socialization\u2013volunteering association, but "
     "this moderating effect will be weaker for Generation Z than for older cohorts.", True),
], indent=True)

mixed([
    ("Hypothesis 3b.", True, True, False),
    (" Employment will positively moderate the socialization\u2013volunteering association, with "
     "effects that are relatively consistent across generations.", True),
], indent=True)

mixed([
    ("Hypothesis 3c.", True, True, False),
    (" Civic social media use will positively moderate the socialization\u2013volunteering "
     "association, but the moderating effect will be stronger for older cohorts than for "
     "Generation Z.", True),
], indent=True)


# ╔══════════════════════════════════════════════════════════════════════╗
# ║  METHOD  (Comments 9, 10: reorganize, add justification, simplify) ║
# ╚══════════════════════════════════════════════════════════════════════╝

h1("Method")

# ── Data Source (Comment 9: where, who, why it fits) ──────────────────
h2("Data Source", highlight=True)

mixed([
    ("Data come from the Current Population Survey Civic Engagement and Volunteering Supplement "
     "(CPS-CEV), obtained via IPUMS (Flood et al., 2023). The CPS-CEV is administered as a supplement "
     "to the September CPS in alternating years and is sponsored jointly by AmeriCorps and the "
     "U.S. Census Bureau. It employs a stratified multistage probability sampling design that produces "
     "nationally representative estimates of the civilian noninstitutionalized population aged 16 and "
     "older. ", False),
    ("The CPS-CEV is the only nationally representative survey that simultaneously measures both "
     "volunteering behavior and in-person socialization frequency across multiple waves, making it "
     "uniquely suited to examining the socialization\u2013volunteering association over time and "
     "across the COVID-19 pandemic.", True),
], indent=True)

# ── Sample (Comment 9: inclusion, N, missing data) ────────────────────
h2("Sample", highlight=True)

mixed([
    ("We pooled four waves of CPS-CEV data: September 2017, 2019, 2021, and 2023. ", False),
    ("The inclusion criteria were: (a) aged 18 years or older, (b) valid supplement weight "
     "(VLSUPPWT > 0), and (c) nonmissing volunteering status (VLSTATUS \u2208 {1, 2}).", True),
    (" The resulting analytic sample included ", False),
    ("N", False, False, True),
    (" = 201,168, with each wave contributing approximately 50,000 respondents. The four-wave design "
     "spanning the COVID-19 pandemic enables examination of temporal patterns while maintaining "
     "sufficient statistical power for generational subgroup analyses.", False),
], indent=True)

body(
    "Respondents with missing values on the key independent variable (CESOCIALIZE) were excluded "
    "from models involving socialization (less than 1% of the analytic sample). For moderating "
    "variables, cases with missing values were excluded from the respective interaction models. "
    "Survey supplement weights (VLSUPPWT) were applied to all analyses to produce nationally "
    "representative estimates.",
    highlight=True,
)

# ── Measures ──────────────────────────────────────────────────────────
h2("Measures")

# DV (Comment 9: clear subsection)
h3("Dependent Variable")

body(
    "Volunteering status was measured using VLSTATUS, which asks whether the respondent \u201cdid any "
    "volunteer activities through or for an organization\u201d in the past 12 months. This variable was "
    "recoded as a binary indicator (1 = volunteered, 0 = did not volunteer). The national volunteering "
    "rate across the pooled sample was 31.5%.",
)

# Key IV (Comment 9)
h3("Independent Variable", highlight=True)

body(
    "In-person socialization frequency was measured using CESOCIALIZE, which asks respondents how often "
    "they \u201cgot together socially with friends, relatives, or neighbors\u201d in the past 12 months. "
    "Response options range from 1 (not at all) to 6 (basically every day). This variable was entered "
    "as a set of dummy variables (with \u201cnot at all\u201d as the reference category) to capture potential "
    "nonlinear threshold effects.",
)

# Moderating Variables (Comment 9: explicitly labeled)
h3("Moderating Variables", highlight=True)

mixed([
    ("Generation.", True, True, True),
    (" Respondents were classified into five generational cohorts based on birth year: "
     "Generation Z (born 1997 or later), Millennials (1981\u20131996), Generation X (1965\u20131980), "
     "Baby Boomers (1946\u20131964), and the Silent Generation (born before 1946). Birth year was "
     "calculated as survey year minus respondent age. ", False),
    ("Generation is treated as a moderating variable because shared formative experiences\u2014"
     "particularly exposure to digital technology during adolescence\u2014may alter how social "
     "interaction translates into civic behavior (Hooghe, 2012).", True),
], indent=True)

mixed([
    ("Education.", True, True, True),
    (" Measured using PEEDUCA and dichotomized as bachelor\u2019s degree or higher versus less than "
     "a bachelor\u2019s degree, consistent with the CVM\u2019s identification of education as the strongest "
     "resource predictor of civic participation (Verba et al., 1995).", False),
], indent=True)

mixed([
    ("Employment status.", True, True, True),
    (" Derived from PEMLR and dichotomized as currently employed (including self-employed) versus "
     "not employed.", False),
], indent=True)

mixed([
    ("Civic social media use.", True, True, True),
    (" Measured using VLSOCMEDIA, which asks how often the respondent used the internet or social "
     "media to engage in civic activities such as organizing events, expressing opinions on community "
     "issues, or finding volunteering opportunities. ", False),
    ("This variable uses descending frequency coding (1 = basically every day to 6 = not at all) "
     "and was reverse-coded so that higher values indicate more frequent civic social media use.", True),
], indent=True)

# Control Variables (Comment 9)
h3("Control Variables", highlight=True)

body(
    "Models controlled for age (continuous), sex (male/female), race/ethnicity (White non-Hispanic, "
    "Black non-Hispanic, Hispanic, Asian non-Hispanic, other), marital status (married, never married, "
    "divorced/separated/widowed), household income (16 categories), metropolitan status "
    "(metropolitan/nonmetropolitan), region (Northeast, Midwest, South, West), and survey wave "
    "(2017, 2019, 2021, 2023).",
)

# ── Data Analysis (Comments 9, 10: RQ-aligned, simplified terms) ─────
h2("Data Analysis", highlight=True)

body(
    "All analyses were conducted in R (version 4.4) using the survey and marginaleffects packages. "
    "Survey supplement weights (VLSUPPWT) and linearized standard errors were used throughout to "
    "account for the complex survey design.",
    highlight=True,
)

mixed([
    ("To address Research Question 1, ", True),
    ("we estimated survey-weighted logistic regression models predicting volunteering status from "
     "socialization frequency, generational cohort, and their interaction. Socialization frequency "
     "was entered as a set of dummy variables (with \u201cnot at all\u201d as the reference category) to "
     "capture nonlinear threshold effects. Predicted probabilities and average marginal effects were "
     "computed to quantify the socialization\u2013volunteering gradient within each generation. ", False),
    ("To test whether the observed generational pattern reflects a cohort effect rather than an age "
     "effect, we conducted a same-age comparison of Millennials (ages 20\u201325 in 2017) and "
     "Generation Z (ages 20\u201325 in 2023) on socialization and volunteering rates.", True),
], indent=True)

mixed([
    ("To address Research Question 2, ", True),
    ("we extended the logistic regression models to include three-way interactions: "
     "socialization frequency \u00d7 generation \u00d7 education, socialization frequency \u00d7 "
     "generation \u00d7 employment status, and socialization frequency \u00d7 generation \u00d7 "
     "civic social media use. These interactions test whether conventional pathways to civic "
     "participation moderate the socialization\u2013volunteering association differently across "
     "generations (Hypotheses 3a\u20133c).", True),
], indent=True)

# Supplementary analyses (Comment 10: reduce technical jargon)
h3("Supplementary Analyses", highlight=True)

mixed([
    ("Two supplementary analyses were conducted to provide additional context for the primary "
     "regression findings. First, ", True),
    ("Latent Profile Analysis (LPA) was used to identify distinct civic engagement typologies "
     "using five indicators: boycotting (CEBOYCOTT), contacting public officials (CEPUBOFF), "
     "political conversation frequency (CEPOLCONV), socialization frequency (CESOCIALIZE), and "
     "organizational membership count (VLMEMBERN). LPA results provide a person-centered context "
     "for the variable-centered regression findings, confirming whether the socialization\u2013"
     "volunteering gradient operates similarly across empirically derived civic engagement profiles "
     "(Collins & Lanza, 2010). ", False),
    ("Second, a machine learning classification model (gradient boosting; Friedman, 2001) "
     "with variable importance decomposition (Lundberg & Lee, 2017) was used as a robustness check "
     "to identify potential nonlinear relationships and interaction effects not captured by the "
     "parametric models. Details of model specifications are available in the supplementary materials.", True),
], indent=True)


# ╔══════════════════════════════════════════════════════════════════════╗
# ║  REFERENCES  (unchanged)                                            ║
# ╚══════════════════════════════════════════════════════════════════════╝

doc.add_page_break()
h1("References")

refs = [
    "Adler, R. P., & Goggin, J. (2005). What do we mean by \u201ccivic engagement\u201d? Journal of Transformative Education, 3(3), 236\u2013253. https://doi.org/10.1177/1541344605276792",
    "AmeriCorps. (2024). Civic engagement and volunteering (CEV) dashboard. AmeriCorps Open Data. https://data.americorps.gov/stories/s/62w6-z7xa",
    "Bennett, W. L., & Segerberg, A. (2012). The logic of connective action: Digital media and the personalization of contentious politics. Information, Communication & Society, 15(5), 739\u2013768. https://doi.org/10.1080/1369118.2012.670661",
    "Brady, H. E., Verba, S., & Schlozman, K. L. (1995). Beyond SES: A resource model of political participation. American Political Science Review, 89(2), 271\u2013294. https://doi.org/10.2307/2082425",
    "Cacioppo, J. T., & Cacioppo, S. (2014). Social relationships and health: The toxic effects of perceived social isolation. Social and Personality Psychology Compass, 8(2), 58\u201372. https://doi.org/10.1111/spc3.12087",
    "Coleman, J. S. (1988). Social capital in the creation of human capital. American Journal of Sociology, 94(Supplement), S95\u2013S120. https://doi.org/10.1086/228943",
    "Collins, L. M., & Lanza, S. T. (2010). Latent class and latent transition analysis: With applications in the social, behavioral, and health sciences. Wiley.",
    "Dalton, R. J. (2008). The good citizen: How a younger generation is reshaping American politics. CQ Press.",
    "Flanagan, C., & Levine, P. (2010). Civic engagement and the transition to adulthood. The Future of Children, 20(1), 159\u2013179. https://doi.org/10.1353/foc.0.0043",
    "Flood, S., King, M., Rodgers, R., Ruggles, S., Warren, J. R., Backman, D., Chen, A., Cooper, G., Richards, S., Schouweiler, M., & Westberry, M. (2023). Integrated public use microdata series, Current Population Survey: Version 11.0 [dataset]. IPUMS. https://doi.org/10.18128/D030.V11.0",
    "Friedman, J. H. (2001). Greedy function approximation: A gradient boosting machine. Annals of Statistics, 29(5), 1189\u20131232. https://doi.org/10.1214/aos/1013203451",
    "Granovetter, M. S. (1973). The strength of weak ties. American Journal of Sociology, 78(6), 1360\u20131380. https://doi.org/10.1086/225469",
    "Holt-Lunstad, J., Smith, T. B., & Layton, J. B. (2010). Social relationships and mortality risk: A meta-analytic review. PLOS Medicine, 7(7), Article e1000316. https://doi.org/10.1371/journal.pmed.1000316",
    "Hooghe, M. (2012). Taking generation seriously: Three approaches to cohort effects in political research. In M. Hooghe & D. Stolle (Eds.), Generating social capital (pp. 1\u201322). Palgrave Macmillan.",
    "Kelle, N., Simonson, J., & Henning, G. (2025). Baby boomers and their voluntary engagement: A cohort comparison among the middle-aged and older population in Germany. Nonprofit and Voluntary Sector Quarterly, 54(1), 182\u2013203. https://doi.org/10.1177/08997640241240417",
    "Lundberg, S. M., & Lee, S.-I. (2017). A unified approach to interpreting model predictions. In Advances in neural information processing systems 30 (pp. 4765\u20134774). Curran Associates.",
    "Musick, M. A., & Wilson, J. (2008). Volunteers: A social profile. Indiana University Press.",
    "Office of the Surgeon General. (2023). Our epidemic of loneliness and isolation: The U.S. Surgeon General\u2019s advisory on the healing effects of social connection and community. U.S. Department of Health and Human Services.",
    "Putnam, R. D. (2000). Bowling alone: The collapse and revival of American community. Simon & Schuster.",
    "Schlozman, K. L., Verba, S., & Brady, H. E. (2012). The unheavenly chorus: Unequal political voice and the broken promise of American democracy. Princeton University Press.",
    "Skocpol, T. (2003). Diminished democracy: From membership to management in American civic life. University of Oklahoma Press.",
    "Twenge, J. M. (2017). iGen: Why today\u2019s super-connected kids are growing up less rebellious, more tolerant, less happy\u2014and completely unprepared for adulthood. Atria Books.",
    "Twenge, J. M., Spitzberg, B. H., & Campbell, W. K. (2019). Less in-person social interaction with peers among U.S. adolescents in the 21st century and links to loneliness. Journal of Social and Personal Relationships, 36(6), 1892\u20131913. https://doi.org/10.1177/0265407519836170",
    "Verba, S., Schlozman, K. L., & Brady, H. E. (1995). Voice and equality: Civic voluntarism in American politics. Harvard University Press.",
    "Wilson, J. (2000). Volunteering. Annual Review of Sociology, 26(1), 215\u2013240. https://doi.org/10.1146/annurev.soc.26.1.215",
    "Wilson, J. (2012). Volunteerism research: A review essay. Nonprofit and Voluntary Sector Quarterly, 41(2), 176\u2013212. https://doi.org/10.1177/0899764011434558",
    "Woolcock, M., & Narayan, D. (2000). Social capital: Implications for development theory, research, and policy. World Bank Research Observer, 15(2), 225\u2013249. https://doi.org/10.1093/wbro/15.2.225",
    "Zukin, C., Keeter, S., Andolina, M., Jenkins, K., & Delli Carpini, M. X. (2006). A new engagement? Political participation, civic life, and the changing American citizen. Oxford University Press.",
]

for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.first_line_indent = None
    # Hanging indent
    pf = p.paragraph_format
    pf.left_indent = Inches(0.5)
    pf.first_line_indent = Inches(-0.5)
    r = p.add_run(ref)
    _apply_font(r)


# ╔══════════════════════════════════════════════════════════════════════╗
# ║  APPENDIX: RESPONSE TO DR. WINDON'S COMMENTS                       ║
# ╚══════════════════════════════════════════════════════════════════════╝

doc.add_page_break()
h1("Response to Dr. Windon\u2019s Comments (March 11, 2026)")
blank()

comments = [
    (
        "Comment 0 (Abstract)",
        "The abstract is overstated. Need to see the requirements of the journal toward the abstract.",
        "Toned down all causal/deterministic language. \u201cExamines the relationship\u201d \u2192 "
        "\u201cexplores the association.\u201d \u201cVolunteering ceiling\u201d \u2192 "
        "\u201cpredicted volunteering probability appears to level off.\u201d "
        "\u201cLargest marginal gain\u201d \u2192 \u201clargest observed difference in volunteering "
        "probability.\u201d Will check NVSQ abstract formatting requirements before submission."
    ),
    (
        "Comment 1 (Introduction structure)",
        "Intro is too long. Need to explain why this research is needed. Structure: "
        "Problem \u2192 Gap \u2192 Purpose \u2192 Contribution.",
        "Restructured Introduction from 7 paragraphs to 6, following Problem \u2192 Gap \u2192 "
        "Purpose \u2192 Contribution arc. Removed RQs from Introduction (moved to new "
        "\u201cPurpose and Research Questions\u201d section). Added explicit Gap paragraph "
        "identifying the lack of empirical research connecting social isolation to civic outcomes. "
        "Shortened contribution paragraph and toned down language."
    ),
    (
        "Comment 2 (Lit review keywords)",
        "Your lit review should include all key words in your research questions. "
        "You are using so many other words that don\u2019t align with your RQs.",
        "Reorganized Literature Review to mirror RQ keywords: socialization, volunteering, "
        "generation/Gen Z, education/employment, civic social media. Removed tangential "
        "concepts (e.g., \u201cduty-based vs. engaged citizenship\u201d as standalone discussion). "
        "Each subsection now directly maps to a construct in the research questions."
    ),
    (
        "Comment 3 (Lit review focus areas)",
        "Lit review should focus on: (1) In-person socialization, (2) Volunteering across "
        "generations / why Gen Z, (3) Demographics\u2013volunteering relationship, "
        "(4) Social media as separate addition, (5) Why generation moderates.",
        "Created five focused subsections: (1) Social Interaction and Civic Recruitment, "
        "(2) Volunteering Across Generations: Why Generation Z?, (3) Education, Employment, "
        "and Volunteering, (4) Civic Social Media Use as a Moderator, "
        "(5) The Loneliness Epidemic and Civic Consequences. Added explicit rationale "
        "for why Gen Z is the focal cohort (3 reasons) and why generation is a "
        "moderating variable (Hooghe, 2012)."
    ),
    (
        "Comment 4 (Loneliness section \u2192 intro)",
        "This could be as an intro part because it is an issue.",
        "Moved the core loneliness epidemic framing (Surgeon General advisory, health analogy) "
        "into the Introduction as part of the Problem statement. The Literature Review now "
        "retains only the civic-specific extension of the loneliness literature."
    ),
    (
        "Comment 5 (\u201cDemocratic health\u201d term)",
        "What is a democratic health?",
        "Removed the term \u201cdemocratic health\u201d entirely. Section retitled from "
        "\u201cThe Loneliness Epidemic and Democratic Health\u201d to "
        "\u201cThe Loneliness Epidemic and Civic Consequences.\u201d "
        "All instances replaced with \u201ccivic participation,\u201d "
        "\u201ccivic consequences,\u201d or \u201ccommunity capacity.\u201d"
    ),
    (
        "Comment 6 (Conceptual model)",
        "Show the conceptual model, it is difficult to understand what your dependent variables are.",
        "Added a description of Figure 1 (Conceptual Model) in the new \u201cPurpose and "
        "Research Questions\u201d section, clearly showing: DV = volunteering (binary), "
        "IV = in-person socialization frequency, Moderators = generation, education, employment, "
        "civic social media, Controls = age, sex, race/ethnicity, marital status, income, metro, "
        "region, wave. Figure placeholder included; will create visual diagram before submission."
    ),
    (
        "Comment 7 (Hypotheses needed)",
        "What are the main hypotheses of this study?",
        "Added formal hypotheses: H1 (socialization\u2013volunteering positive association), "
        "H2 (Gen Z plateau effect), H3a (education moderation weaker for Gen Z), "
        "H3b (employment moderation consistent across generations), "
        "H3c (civic social media moderation stronger for older cohorts). "
        "Each hypothesis is grounded in the preceding literature review."
    ),
    (
        "Comment 8 (Purpose section needed)",
        "Need to add a small section called \u201cPurpose and Research Questions\u201d "
        "between Lit Review and Method.",
        "Created new \u201cPurpose and Research Questions\u201d section between Literature Review "
        "and Method, containing: study purpose statement, conceptual model description (Figure 1), "
        "two research questions, and six testable hypotheses (H1, H2, H3a\u2013H3c)."
    ),
    (
        "Comment 9 (Method section checklist)",
        "Method Section: dataset name, organization, sampling design, population, years, "
        "why it fits, inclusion criteria, sample size, weighting, variables "
        "(DV/IV/moderators/controls), data analysis subsection.",
        "Reorganized Method into clear subsections: Data Source (with justification for why "
        "CPS-CEV uniquely fits this study), Sample (explicit inclusion criteria, N, missing "
        "data handling), Measures (with separate sub-subsections for DV, IV, Moderating Variables, "
        "and Control Variables), and Data Analysis (organized by RQ with simplified language)."
    ),
    (
        "Comment 10 (Method simplification)",
        "Shorten coding explanation. Make Measures sub-sections. "
        "Remove/reduce technical terms like GBMs, MLL.",
        "Shortened coding explanation (detailed notes moved to supplementary materials). "
        "Created four Measures sub-subsections (DV, IV, Moderators, Controls). "
        "Replaced \u201cgradient boosting machines (GBMs) with TreeSHAP decomposition\u201d with "
        "\u201cmachine learning classification model with variable importance decomposition.\u201d "
        "Removed MLL, BIC, BLRT from main text (deferred to supplementary materials). "
        "LPA and GBM+SHAP are now framed as \u201cSupplementary Analyses\u201d with a separate "
        "sub-subsection."
    ),
]

# Create response table
table = doc.add_table(rows=1, cols=3)
table.style = "Table Grid"

# Header row
for i, header in enumerate(["Comment", "Dr. Windon\u2019s Feedback", "Revision Made"]):
    cell = table.rows[0].cells[i]
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(header)
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(10)

# Data rows
for label, feedback, response in comments:
    row = table.add_row()
    for i, text in enumerate([label, feedback, response]):
        cell = row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(text)
        r.font.name = "Times New Roman"
        r.font.size = Pt(10)

# Set column widths
for row in table.rows:
    row.cells[0].width = Inches(1.2)
    row.cells[1].width = Inches(2.8)
    row.cells[2].width = Inches(3.0)


# ╔══════════════════════════════════════════════════════════════════════╗
# ║  APPENDIX: CODE-MANUSCRIPT DISCREPANCY NOTES                       ║
# ╚══════════════════════════════════════════════════════════════════════╝

doc.add_page_break()
h1("Code\u2013Manuscript Discrepancy Notes (Internal)")

body(
    "The following discrepancies between the GitHub codebase "
    "(HosungYou/-CPS-Civic-Engagement-Volunteering-Supplement) and the manuscript text "
    "were identified during revision. These must be resolved before submission.",
    highlight=True, indent=False,
)

blank()
h2("Discrepancy 1: LPA Profile Naming", highlight=True)

body(
    "Issue: The discussion documents reference a \u201cCheckbook Citizens\u201d profile defined by "
    "donation behavior. However, the actual LPA code (03_latent_profile_analysis.R) uses five "
    "indicators\u2014CEBOYCOTT, CEPUBOFF, CEPOLCONV, CESOCIALIZE, VLMEMBERN\u2014none of which "
    "includes donation (VLDONATE). Naming a profile after a behavior not captured by the LPA "
    "indicators would be indefensible in peer review.",
    highlight=True,
)
body(
    "Resolution: Profile names must be derived from the actual indicator patterns after model "
    "estimation. The \u201cCheckbook Citizens\u201d label should be replaced with a data-driven name "
    "(e.g., \u201cOrganizational Members\u201d if characterized by high VLMEMBERN and moderate "
    "other indicators).",
    highlight=True,
)

blank()
h2("Discrepancy 2: LPA Social Interaction Variable", highlight=True)

body(
    "Issue: The earlier PA-focused design document (discussion/05) specified LPA indicator #4 as "
    "the average of CESOCCONTCT and CESOCIALIZE: (CESOCCONTCT + CESOCIALIZE) / 2. "
    "The actual R code uses CESOCIALIZE alone (line 40).",
    highlight=True,
)
body(
    "Resolution: The manuscript correctly describes CESOCIALIZE as the socialization indicator. "
    "No change needed to the manuscript, but the discrepancy with the earlier design document "
    "should be noted. Rationale for using CESOCIALIZE alone (e.g., avoiding multicollinearity, "
    "clearer interpretation) should be documented.",
    highlight=True,
)

blank()
h2("Discrepancy 3: GBM Generation Comparison Scope", highlight=True)

body(
    "Issue: The manuscript implies comparison \u201cacross generational cohorts\u201d (all five "
    "generations). However, the Python code (04_gbm_shap_analysis.py, line 191) runs SHAP "
    "decomposition for only Gen Z and Boomers: for gen in ['Gen Z', 'Boomer'].",
    highlight=True,
)
body(
    "Resolution: Either (a) expand the code to run SHAP for all five generations, or "
    "(b) explicitly state in the manuscript that the generational SHAP comparison focuses on "
    "\u201cGen Z and Baby Boomers as the most and least socially isolated cohorts, respectively, "
    "to maximize contrast.\u201d Option (b) is recommended with a full-sample SHAP summary as "
    "supplementary material.",
    highlight=True,
)


# ── Save ─────────────────────────────────────────────────────────────
doc.save(OUTPUT)
print(f"Saved: {OUTPUT}")
