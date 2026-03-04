#!/usr/bin/env python3
"""
Generate the research proposal Word document with proper APA 7th formatting.
Bowling Alone, Scrolling Together: Social Isolation and the Generational Fracture
of Civic Engagement in America
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ── Page setup ──
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# ── Style defaults ──
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
paragraph_format = style.paragraph_format
paragraph_format.space_after = Pt(0)
paragraph_format.space_before = Pt(0)
paragraph_format.line_spacing = 2.0

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def add_para(text, bold=False, italic=False, alignment=None, space_after=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = bold
    run.italic = italic
    if alignment:
        p.alignment = alignment
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    return p

def add_rich_para(segments, alignment=None, space_after=None):
    """segments: list of (text, bold, italic) tuples"""
    p = doc.add_paragraph()
    for text, bold, italic in segments:
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.bold = bold
        run.italic = italic
    if alignment:
        p.alignment = alignment
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    return p

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)
    # Data
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            run = cell.paragraphs[0].add_run(str(val))
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
    doc.add_paragraph()  # spacing
    return table

# ═══════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
add_para('Bowling Alone, Scrolling Together:', bold=True,
         alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
add_para('Social Isolation and the Generational Fracture of Civic Engagement in America',
         bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
add_para('Research Proposal and Exploratory Findings',
         italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
add_para('Hosung You and Suzanna Windon',
         alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
add_para('Department of Agricultural Economics, Sociology, and Education',
         alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
add_para('The Pennsylvania State University',
         alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
add_para('March 2026', alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
add_para('Data Source: Current Population Survey \u2014 Civic Engagement and Volunteering Supplement (CPS-CEV), 2017, 2019, 2021, 2023',
         italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)

doc.add_page_break()

# ═══════════════════════════════════════════════════════
# 1. INTRODUCTION
# ═══════════════════════════════════════════════════════
add_heading_styled('1. Introduction', level=1)

add_para(
    'While exploring the CPS-CEV data for our civic engagement project, we encountered a finding '
    'that reframes the entire study. The data reveals a phenomenon far more consequential than '
    'state-level variation in volunteering patterns: social isolation is the strongest behavioral '
    'correlate of civic disengagement across all generations, and Generation Z stands at the '
    'epicenter of a socialization crisis that education, employment, and digital connectivity '
    'fail to resolve.'
)

add_para(
    'This document presents exploratory descriptive analyses conducted to understand the data '
    'structure and identify promising research directions. The statistics reported here are based '
    'on cross-tabulations and frequency distributions; the formal analytical methods (Latent '
    'Profile Analysis, Gradient Boosting Machines with SHAP decomposition) differ substantially '
    'and are specified in the Methods section.'
)

# ═══════════════════════════════════════════════════════
# 2. THEORETICAL FOUNDATIONS
# ═══════════════════════════════════════════════════════
add_heading_styled('2. Theoretical Foundations', level=1)

add_heading_styled('2.1 Social Capital and Democratic Participation', level=2)

add_para(
    'The relationship between social connectedness and civic participation has deep theoretical '
    'roots. Putnam (2000) documented the broad decline of associational life in late 20th-century '
    'America, arguing that weakening social networks eroded both bonding capital (solidarity within '
    'groups) and bridging capital (connections across diverse groups). His framework predicts that '
    'declines in face-to-face social interaction should produce downstream reductions in civic '
    'engagement\u2014a prediction our data strongly supports.'
)
add_para(
    'Coleman (1988) established the theoretical link between social capital and human capital '
    'formation, arguing that social networks create obligations, expectations, and information '
    'channels that facilitate collective action. Granovetter\'s (1973) insight that "weak ties" '
    'often matter more than strong ties for accessing new information and opportunities is '
    'particularly relevant: the transition from zero social contact to even minimal contact '
    'activates weak-tie networks that may channel civic recruitment.'
)

add_heading_styled('2.2 The Civic Voluntarism Model', level=2)

add_para(
    'Verba et al. (1995) identified three categories of factors predicting civic participation: '
    'resources (time, money, civic skills), engagement (political interest, civic duty), and '
    'recruitment (being asked to participate through social networks). Our finding that '
    'socialization frequency is the strongest correlate of volunteering across all generations '
    'aligns with the recruitment mechanism: social interaction creates the occasions through '
    'which individuals are recruited into civic activities. Brady et al. (1995) further '
    'demonstrated that civic skills acquired through organizational involvement mediate the '
    'education-participation relationship\u2014a mechanism that requires social connection as a '
    'prerequisite.'
)

add_heading_styled('2.3 The Loneliness Epidemic and Democratic Health', level=2)

add_para(
    'The U.S. Surgeon General\'s 2023 advisory declared loneliness and social isolation a public '
    'health epidemic, documenting extensive consequences for mental and physical health (Office '
    'of the Surgeon General, 2023). However, the advisory focused almost exclusively on health '
    'outcomes. The civic consequences of the loneliness epidemic\u2014its implications for democratic '
    'participation, community capacity, and collective governance\u2014remain largely unexamined. '
    'Our study addresses this gap by connecting nationally representative data on social isolation '
    'to measurable civic engagement outcomes.'
)

add_heading_styled('2.4 Generational Shifts in Civic Engagement', level=2)

add_para(
    'Zukin et al. (2006) documented generational differences in civic participation styles, '
    'finding that younger cohorts favored volunteering and community service over electoral '
    'participation. Skocpol (2003) traced the broader shift from membership-based civic '
    'organizations to professionally managed advocacy groups. Twenge (2017) identified '
    'smartphone adoption as a potential inflection point in adolescent social behavior, with '
    'sharp declines in face-to-face socializing beginning around 2012. Bennett and Segerberg '
    '(2012) theorized the rise of "connective action"\u2014individualized, digitally mediated civic '
    'expression that operates through personal sharing rather than organizational '
    'mobilization\u2014as a complement or substitute for traditional collective action.'
)
add_para(
    'Our exploratory analysis integrates these perspectives by examining how social isolation, '
    'generational membership, and digital connectivity jointly shape civic engagement patterns '
    'in a nationally representative sample spanning the COVID-19 pandemic.'
)

# ═══════════════════════════════════════════════════════
# 3. DATA AND METHODS OVERVIEW
# ═══════════════════════════════════════════════════════
add_heading_styled('3. Data and Methods Overview', level=1)

add_heading_styled('3.1 Data Source and Sample', level=2)

add_para(
    'Data come from the Current Population Survey\u2014Civic Engagement and Volunteering Supplement '
    '(CPS-CEV), obtained via IPUMS (Flood et al., 2023). The CPS-CEV is administered as a '
    'supplement to the September CPS in alternating years. We use four waves: 2017, 2019, 2021, '
    'and 2023. The analytic sample includes all U.S. adults aged 18 and older with valid '
    'supplement weights (VLSUPPWT > 0) and non-missing volunteering status (VLSTATUS \u2208 {1, 2}), '
    'yielding N = 201,168.'
)

add_para('Table 1. Variable Coding (Verified)', bold=True, space_after=6)
add_table(
    ['Variable', 'Type', 'Coding Direction', 'Verified By'],
    [
        ['VLSTATUS', 'Binary', '1 = Volunteered, 2 = Did not', 'Education cross-validation'],
        ['VLDONATE', 'Binary', '1 = Did not donate, 2 = Donated', 'Direct verification'],
        ['CEBOYCOTT', 'Binary', '1 = No, 2 = Yes', 'Direct verification'],
        ['CEPUBOFF', 'Binary', '1 = No, 2 = Yes', 'Direct verification'],
        ['CEPOLCONV', '6-point freq', '1 = Never \u2192 6 = Daily', 'Education cross-validation'],
        ['CESOCCONTCT', '6-point freq', '1 = Never \u2192 6 = Daily', 'Education cross-validation'],
        ['CESOCIALIZE', '6-point freq', '1 = Never \u2192 6 = Daily', 'Education cross-validation'],
        ['VLSOCMEDIA', '6-point freq', '1 = Daily \u2192 6 = Never*', 'Distribution analysis'],
    ]
)
add_para(
    '*Note: CE* variables (Civic Engagement supplement) use ascending frequency; VL* variables '
    '(Volunteering supplement) use descending frequency. No reverse-coding of CE* variables is needed.',
    italic=True
)

add_heading_styled('3.2 Descriptive Exploration vs. Formal Analysis', level=2)

add_para(
    'The findings presented in Sections 4 through 6 are based on descriptive cross-tabulations '
    'and frequency distributions designed to reveal data structure and generate hypotheses. These '
    'exploratory statistics do not constitute the formal analytical methods of the proposed study. '
    'The formal analysis will employ Latent Profile Analysis (LPA) for typological discovery '
    '(Collins & Lanza, 2010; Muth\u00e9n & Muth\u00e9n, 2000) and Gradient Boosting Machines with '
    'TreeSHAP decomposition (Lundberg & Lee, 2017) for interpretable prediction, as detailed in '
    'Section 7.'
)

add_heading_styled('3.3 Generational Definitions', level=2)

add_para('Table 2. Generational Cohort Definitions', bold=True, space_after=6)
add_table(
    ['Generation', 'Birth Years', 'Age Range (2023)'],
    [
        ['Gen Z', '\u2265 1997', '18\u201326'],
        ['Millennial', '1981\u20131996', '27\u201342'],
        ['Gen X', '1965\u20131980', '43\u201358'],
        ['Baby Boomer', '1946\u20131964', '59\u201377'],
        ['Silent', '< 1946', '\u2265 78'],
    ]
)

# ═══════════════════════════════════════════════════════
# 4. EXPLORATORY FINDINGS
# ═══════════════════════════════════════════════════════
add_heading_styled('4. Exploratory Findings: The Socialization\u2013Civic Engagement Nexus', level=1)

add_heading_styled('4.1 The Socialization Crisis Is Generational, Not Pandemic', level=2)

add_para(
    'The CPS-CEV supplement asks respondents how often they "got together socially with friends, '
    'relatives, or neighbors" (CESOCIALIZE). The percentage reporting "not at all" reveals a '
    'striking generational divide that predates the COVID-19 pandemic.'
)

add_para('Table 3. Percentage Never Socializing, by Generation and Year', bold=True, space_after=6)
add_table(
    ['Generation', '2017', '2019', '2021', '2023'],
    [
        ['Gen Z', '44.6%', '46.3%', '48.4%', '46.6%'],
        ['Millennial', '30.9%', '32.8%', '32.7%', '29.6%'],
        ['Gen X', '20.8%', '24.9%', '25.6%', '24.0%'],
        ['Boomer', '17.3%', '19.7%', '20.9%', '19.2%'],
        ['Silent', '17.2%', '20.2%', '20.7%', '22.4%'],
    ]
)

add_para(
    'Nearly half of Gen Z reports never socializing in person. This was true in 2017\u2014two years '
    'before COVID-19. The pandemic modestly worsened the pattern (48.4% in 2021) but did not '
    'create it, and recovery has been minimal.'
)

add_heading_styled('4.2 Cohort Effect, Not Age Effect', level=2)

add_para(
    'To rule out the possibility that young adults are simply always more isolated, we compared '
    'Millennials and Gen Z at the same age.'
)

add_para('Table 4. Same-Age Comparison: Millennials (2017) vs. Gen Z (2023) at Ages 20\u201325',
         bold=True, space_after=6)
add_table(
    ['Metric', 'Millennials at 20\u201325 (2017)', 'Gen Z at 20\u201325 (2023)', 'Difference'],
    [
        ['Never socialize', '40.0%', '47.4%', '+7.4pp'],
        ['Volunteering rate', '26.3%', '22.4%', '\u22123.9pp'],
        ['Civically disengaged', '53.1%', '56.6%', '+3.5pp'],
        ['BA+ education', '25.9%', '25.9%', '0.0pp'],
    ]
)

add_para(
    'Same age, identical education profile, but Gen Z is 7.4 percentage points more socially '
    'isolated and 3.5 points more civically disengaged. This is a generational shift, not a '
    'life-course artifact.'
)

add_heading_styled('4.3 Education and Employment Do Not Resolve the Crisis', level=2)

add_para('Table 5. Never-Socializing Rates Among College Graduates, by Generation',
         bold=True, space_after=6)
add_table(
    ['Generation', 'BA+ Never-Socialize Rate', 'BA+ Volunteering Rate'],
    [
        ['Gen Z', '42.1%', '34.0%'],
        ['Millennial', '25.5%', '43.1%'],
        ['Gen X', '16.5%', '52.5%'],
        ['Boomer', '13.4%', '47.1%'],
    ]
)

add_para(
    'College-educated Gen Z members are three times more likely to never socialize than '
    'college-educated Boomers (42.1% vs. 13.4%). A bachelor\'s degree does not resolve Gen Z\'s '
    'socialization deficit. Employment shows a similarly negligible effect: 45.6% of employed '
    'Gen Z never socializes, compared to 48.3% of non-employed Gen Z\u2014a mere 2.7 percentage '
    'point difference. For Boomers, the comparable gap is only 0.9pp (18.6% vs. 19.5%). This '
    'distinguishes the current crisis from historical patterns where workplace integration served '
    'as a pathway to social connection and civic recruitment (Verba et al., 1995).'
)

add_heading_styled('4.4 The "First Step" Effect: Nonlinear Socialization\u2013Volunteering Gradient', level=2)

add_para('Table 6. Volunteering Rate by Socialization Frequency and Generation',
         bold=True, space_after=6)
add_table(
    ['Socialization Level', 'Gen Z', 'Millennial', 'Gen X', 'Boomer'],
    [
        ['Never', '15.7%', '19.1%', '18.7%', '13.7%'],
        ['A few times/year', '25.6%', '30.1%', '32.3%', '25.4%'],
        ['Once a month', '32.4%', '37.2%', '40.0%', '33.3%'],
        ['A few times/month', '33.4%', '38.6%', '43.0%', '36.8%'],
        ['A few times/week', '31.9%', '42.0%', '48.7%', '41.5%'],
        ['Daily', '34.6%', '42.5%', '44.3%', '40.7%'],
    ]
)

add_para(
    'The single largest jump in volunteering occurs at the transition from "never" to "a few '
    'times a year"\u2014a gain of 10 to 14 percentage points across all generations. This is the '
    'most efficient intervention point: moving someone from complete social isolation to even '
    'minimal social contact produces the largest marginal increase in civic participation.'
)
add_para(
    'However, Gen Z exhibits a ceiling effect at approximately 34% volunteering\u2014regardless '
    'of how frequently they socialize. Other generations continue to see gains at higher '
    'socialization frequencies, but Gen Z\'s volunteering rate plateaus. This suggests a '
    'structural barrier beyond socialization alone.'
)

add_heading_styled('4.5 Social Isolation Creates a Universal Civic Floor', level=2)

add_para('Table 7. Volunteering Rates by Education \u00d7 Socialization \u00d7 Generation',
         bold=True, space_after=6)
add_table(
    ['Condition', 'Gen Z', 'Millennial', 'Boomer'],
    [
        ['HS diploma + Isolated', '12.2%', '11.5%', '~12%'],
        ['HS diploma + Connected', '26.3%', '25.5%', '25.4%'],
        ['BA+ degree + Isolated', '27.7%', '33.4%', '29.5%'],
        ['BA+ degree + Connected', '42.8%', '51.3%', '54.5%'],
    ]
)

add_para(
    'The civic "floor"\u2014the volunteering rate of socially isolated, high-school-educated '
    'adults\u2014is approximately 12% regardless of generation. Social isolation creates a universal '
    'civic minimum. The generational differences emerge at the civic "ceiling," where Gen Z\'s '
    'maximum (42.8%) lags behind older cohorts (51\u201355%).'
)

add_heading_styled('4.6 The Counterintuitive Finding', level=2)

add_para(
    'Perhaps the most striking result: the most socially connected, college-educated, politically '
    'conversant Gen Z members (N = 246) volunteer at 45.5%\u2014higher than the average Baby Boomer '
    '(32.3%). The "generational decline" narrative collapses once socialization is accounted for. '
    'The crisis is not that a generation has abandoned civic life; it is that nearly half a '
    'generation has been cut off from the social infrastructure that makes civic life possible.'
)

# ═══════════════════════════════════════════════════════
# 5. THE DIGITAL NATIVE PARADOX
# ═══════════════════════════════════════════════════════
add_heading_styled('5. The Digital Native Paradox', level=1)

add_heading_styled('5.1 Social Media Compensates for Isolation\u2014But More for Boomers', level=2)

add_para('Table 8. Social Media Compensation Effect Among Socially Isolated Individuals',
         bold=True, space_after=6)
add_table(
    ['Generation', 'Isolated + SM Use', 'Isolated + No SM', 'Compensation Effect'],
    [
        ['Gen Z', '22.6%', '17.0%', '+5.5pp'],
        ['Millennial', '28.6%', '21.5%', '+7.1pp'],
        ['Gen X', '29.5%', '22.7%', '+6.8pp'],
        ['Boomer', '28.3%', '17.0%', '+11.3pp'],
        ['Silent', '23.2%', '12.2%', '+10.9pp'],
    ]
)

add_para(
    'Among socially connected individuals, the pattern is even starker: for Gen Z, civic social '
    'media use adds only 1.6pp to volunteering, while for connected Boomers it adds 9.9pp. '
    'Social media appears to function as a genuinely new civic channel for older adults but is '
    'merely ambient noise for a generation that has never known life without it.'
)

add_heading_styled('5.2 The Civic Substitution: From Volunteering to Boycotting', level=2)

add_para('Table 9. Change in Civic Activities, 2017\u20132023, by Generation',
         bold=True, space_after=6)
add_table(
    ['Generation', 'Volunteering \u0394', 'Boycotting \u0394', 'Donating \u0394', 'Contacting Officials \u0394'],
    [
        ['Gen Z', '\u22122.6pp', '+8.3pp', '+7.1pp', '+0.6pp'],
        ['Millennial', '+2.8pp', '+6.4pp', '+4.4pp', '+0.1pp'],
        ['Gen X', '\u22124.3pp', '+4.9pp', '\u22122.9pp', '\u22122.1pp'],
        ['Boomer', '\u22122.3pp', '+4.0pp', '\u22121.9pp', '\u22124.0pp'],
    ]
)

add_para(
    'Gen Z\'s boycotting rate more than doubled from 6.5% to 14.8% between 2017 and 2023, even '
    'as volunteering declined. This suggests a "consumerization" of civic engagement\u2014replacing '
    'time-intensive, community-embedded participation with individual, market-mediated action. '
    'The phenomenon aligns with Bennett and Segerberg\'s (2012) concept of connective action '
    'displacing traditional collective action.'
)

# ═══════════════════════════════════════════════════════
# 6. COVID AS AMPLIFIER, NOT CAUSE
# ═══════════════════════════════════════════════════════
add_heading_styled('6. COVID-19 as Amplifier, Not Cause', level=1)

add_para(
    'The four-wave design (2017, 2019, 2021, 2023) allows examination of three distinct '
    'generational responses to the pandemic:'
)

add_rich_para([
    ('Gen Z: Deepening isolation with civic substitution. ', True, False),
    ('Social isolation peaked during COVID but was already severe pre-pandemic. Civic '
     'disengagement actually improved slightly by 2023\u2014but only because boycotting doubled, '
     'not because volunteering or formal participation recovered.', False, False),
])

add_rich_para([
    ('Millennials: Steady civic maturation despite disruption. ', True, False),
    ('The only generation showing continuous decline in disengagement and continuous growth in '
     'multi-domain civic activity, even through the pandemic. This may reflect life-course effects '
     'combined with political activation around social justice movements.', False, False),
])

add_rich_para([
    ('Boomers: Classic V-shaped recovery. ', True, False),
    ('Volunteering dropped sharply during COVID (34.2% \u2192 28.4%) but substantially recovered '
     '(31.7%), while donation and boycotting rates remained stable or grew. Their civic '
     'infrastructure proved resilient.', False, False),
])

# ═══════════════════════════════════════════════════════
# 7. PROPOSED RESEARCH FRAMEWORK
# ═══════════════════════════════════════════════════════
add_heading_styled('7. Proposed Research Framework', level=1)

add_heading_styled('7.1 Research Questions', level=2)

add_rich_para([
    ('RQ1 (Typological Structure): ', True, False),
    ('What distinct civic engagement profiles emerge from Latent Profile Analysis of the national '
     'CPS-CEV sample, and how does generational membership predict profile classification?', False, False),
])
add_para(
    'Expected: 4\u20136 profiles including "Isolated Disengaged" (~22%), "Connected Disengaged" '
    '(~14%), "Checkbook Citizens" (~22%), "Consumer Activists" (~8%), "Traditional Volunteers" '
    '(~14%), and "All-around Civic" (~14%). Gen Z will be disproportionately concentrated in '
    'the Isolated Disengaged profile.',
    italic=True
)

add_rich_para([
    ('RQ2 (The Socialization Gradient): ', True, False),
    ('How does the predictive relationship between socialization frequency and volunteering vary '
     'across generational cohorts, and what role do education and digital connectivity play in '
     'moderating this relationship?', False, False),
])
add_para(
    'Expected: GBM + SHAP analysis will identify socialization frequency as the top-ranked '
    'predictor across all profiles, with a pronounced nonlinear "first step" effect.',
    italic=True
)

add_rich_para([
    ('RQ3 (Pandemic Structural Change): ', True, False),
    ('Did the COVID-19 pandemic alter the typological structure of civic engagement, and were '
     'these shifts consistent across generational cohorts?', False, False),
])
add_para(
    'Expected: Profile distribution shifts showing Gen Z\'s "Isolated Disengaged" proportion '
    'peaking in 2021, with counter-trend Millennial growth in active profiles.',
    italic=True
)

add_rich_para([
    ('RQ4 (The Digital Native Paradox): ', True, False),
    ('Does civic use of social media compensate for the civic consequences of social isolation, '
     'and does this compensation vary by generation?', False, False),
])
add_para(
    'Expected: Social media compensation effect approximately +5.5pp for isolated Gen Z versus '
    '+11.3pp for isolated Boomers.',
    italic=True
)

add_heading_styled('7.2 Analytical Framework', level=2)

add_rich_para([
    ('Phase 1: Latent Profile Analysis. ', True, False),
    ('Five indicators: CEBOYCOTT (binary), CEPUBOFF (binary), CEPOLCONV (frequency 1\u20136), '
     'CESOCIALIZE (frequency 1\u20136), VLMEMBERN (count). National sample (N \u2248 200,000). '
     'Model selection via BIC, entropy, and BLRT. Generation as auxiliary variable (BCH method).',
     False, False),
])

add_rich_para([
    ('Phase 2: Profile-Stratified Prediction (GBM + SHAP). ', True, False),
    ('Gradient Boosting Machines predicting volunteering within each profile. TreeSHAP '
     'decomposition of feature contributions. Key predictors: generation, socialization frequency, '
     'education, employment, social media use, geography.', False, False),
])

add_rich_para([
    ('Phase 3: Temporal Comparison. ', True, False),
    ('Profile distribution comparison across 2017, 2019, 2021, 2023 waves. Generation \u00d7 wave '
     'interaction analysis. Pre-pandemic baseline vs. pandemic/post-pandemic periods.', False, False),
])

add_rich_para([
    ('Phase 4: Social Media Moderation Analysis. ', True, False),
    ('Socialization \u00d7 social media \u00d7 generation three-way interaction. SHAP-based decomposition '
     'of compensation effects by profile and generation.', False, False),
])

# ═══════════════════════════════════════════════════════
# 8. SIGNIFICANCE
# ═══════════════════════════════════════════════════════
add_heading_styled('8. Significance and Implications', level=1)

add_heading_styled('8.1 For Community Development and Extension', level=2)
add_para(
    'The finding that the "first step" from never socializing to even minimal social contact '
    'produces the largest civic engagement gain (10\u201314pp) has direct programmatic implications. '
    'Community development practitioners and Extension educators need not aim for intensive '
    'social integration\u2014even low-threshold social programming (community meals, drop-in events, '
    'casual gathering spaces) may produce disproportionate civic returns, particularly for '
    'socially isolated young adults.'
)

add_heading_styled('8.2 For Policy', level=2)
add_para(
    'The failure of education and employment to resolve Gen Z\'s socialization deficit challenges '
    'two dominant policy assumptions: (1) that expanding educational access inherently strengthens '
    'civic capacity, and (2) that labor market integration provides a sufficient pathway to '
    'community connection. Our data suggests that civic infrastructure\u2014the physical and '
    'institutional spaces where people encounter one another\u2014requires deliberate investment '
    'independent of human capital development.'
)

add_heading_styled('8.3 For the Loneliness Conversation', level=2)
add_para(
    'By connecting the Surgeon General\'s loneliness epidemic to measurable civic outcomes using '
    'a nationally representative dataset, this study adds democratic health to the growing case '
    'for treating social isolation as a public priority. The finding that approximately 60% of '
    'Gen Z is completely civically disengaged\u2014and that this is primarily a socialization problem, '
    'not a values or attitude problem\u2014reframes the generational discourse from "young people '
    'don\'t care" to "young people aren\'t connected."'
)

# ═══════════════════════════════════════════════════════
# 9. OPEN QUESTIONS
# ═══════════════════════════════════════════════════════
add_heading_styled('9. Open Questions for Discussion', level=1)

questions = [
    ('Scope: ', 'Do we commit to the national framing, or maintain a PA focus with national '
     'context? Our recommendation is national with PA as an optional illustrative case.'),
    ('Title and framing: ', '"Bowling Alone, Scrolling Together" positions us in conversation '
     'with Putnam\'s foundational work. Is this appropriately ambitious?'),
    ('Causal language: ', 'The findings are cross-sectional and associational. Should we frame '
     'this as a predictive-interpretive study (SHAP-based) or pursue quasi-causal methods?'),
    ('The Gen Z ceiling effect: ', 'Gen Z\'s volunteering plateaus at ~34% regardless of '
     'socialization frequency. Should we explicitly investigate this structural barrier?'),
    ('Target journal: ', 'The generational + loneliness + civic engagement combination could '
     'fit American Behavioral Scientist, Social Forces, Nonprofit and Voluntary Sector Quarterly, '
     'or Journal of Community Psychology.'),
    ('Timeline: ', 'If we agree on the national framing, formal LPA analysis can begin '
     'immediately. The data is clean and coding is verified.'),
]
for i, (label, text) in enumerate(questions, 1):
    add_rich_para([
        (f'{i}. {label}', True, False),
        (text, False, False),
    ])

# ═══════════════════════════════════════════════════════
# REFERENCES (APA 7th)
# ═══════════════════════════════════════════════════════
add_heading_styled('References', level=1)

references = [
    'Bennett, W. L., & Segerberg, A. (2012). The logic of connective action: Digital media and the personalization of contentious politics. Information, Communication & Society, 15(5), 739\u2013768.',
    'Brady, H. E., Verba, S., & Schlozman, K. L. (1995). Beyond SES: A resource model of political participation. American Political Science Review, 89(2), 271\u2013294.',
    'Coleman, J. S. (1988). Social capital in the creation of human capital. American Journal of Sociology, 94(Supplement), S95\u2013S120.',
    'Collins, L. M., & Lanza, S. T. (2010). Latent class and latent transition analysis: With applications in the social, behavioral, and health sciences. Wiley.',
    'Flood, S., King, M., Rodgers, R., Ruggles, S., Warren, J. R., Backman, D., Chen, A., Cooper, G., Richards, S., Schouweiler, M., & Westberry, M. (2023). Integrated public use microdata series, Current Population Survey: Version 11.0 [dataset]. IPUMS. https://doi.org/10.18128/D030.V11.0',
    'Granovetter, M. S. (1973). The strength of weak ties. American Journal of Sociology, 78(6), 1360\u20131380.',
    'Lundberg, S. M., & Lee, S.-I. (2017). A unified approach to interpreting model predictions. In I. Guyon et al. (Eds.), Advances in neural information processing systems 30 (pp. 4765\u20134774). Curran Associates.',
    'Muth\u00e9n, B., & Muth\u00e9n, L. K. (2000). Integrating person-centered and variable-centered analyses: Growth mixture modeling with latent trajectory classes. Alcoholism: Clinical and Experimental Research, 24(6), 882\u2013891.',
    'Office of the Surgeon General. (2023). Our epidemic of loneliness and isolation: The U.S. Surgeon General\u2019s advisory on the healing effects of social connection and community. U.S. Department of Health and Human Services.',
    'Putnam, R. D. (2000). Bowling alone: The collapse and revival of American community. Simon & Schuster.',
    'Rosenstone, S. J., & Hansen, J. M. (1993). Mobilization, participation, and democracy in America. Macmillan.',
    'Skocpol, T. (2003). Diminished democracy: From membership to management in American civic life. University of Oklahoma Press.',
    'Twenge, J. M. (2017). iGen: Why today\u2019s super-connected kids are growing up less rebellious, more tolerant, less happy\u2014and completely unprepared for adulthood. Atria Books.',
    'Verba, S., Schlozman, K. L., & Brady, H. E. (1995). Voice and equality: Civic voluntarism in American politics. Harvard University Press.',
    'Zukin, C., Keeter, S., Andolina, M., Jenkins, K., & Delli Carpini, M. X. (2006). A new engagement? Political participation, civic life, and the changing American citizen. Oxford University Press.',
]

for ref in references:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.5)
    run = p.add_run(ref)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

# ── Closing note ──
doc.add_paragraph()
add_para(
    'All statistics in this document are based on exploratory descriptive analysis of the CPS-CEV '
    'supplement (N = 201,168 adults with valid supplement weights across four waves). Formal '
    'analysis with survey weights, model selection procedures, and inferential testing will '
    'follow upon agreement on research design.',
    italic=True
)

# ── Save ──
out_path = os.path.join(os.path.dirname(__file__),
                        'Bowling_Alone_Scrolling_Together_Proposal.docx')
doc.save(out_path)
print(f'Saved to {out_path}')
