#!/usr/bin/env python3
"""
Build APA-style DOCX manuscript from markdown source files.

Output: Bowling_Alone_CLEAN.docx

CRITICAL: Each heading is immediately followed by its content paragraphs
before the next heading appears (no heading clustering bug).
"""

import re
import os
import csv
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn

BASE_DIR = Path(__file__).parent.parent  # project root


# ---------------------------------------------------------------------------
# Helper functions
# ---------------------------------------------------------------------------

def set_run_font(run, name="Times New Roman", size=12, bold=False, italic=False):
    """Set font properties on a run."""
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    # Force Times New Roman for East Asian fallback
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = run._element.makeelement(qn("w:rFonts"), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:cs"), name)


def set_paragraph_spacing(paragraph, before=0, after=0, line_spacing=2.0):
    """Set paragraph spacing (APA double-space)."""
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line_spacing


def add_heading_apa(doc, text, level=1):
    """Add an APA-style heading.
    Level 1: Centered, Bold, Title Case
    Level 2: Flush Left, Bold, Title Case
    Level 3: Flush Left, Bold Italic, Title Case
    """
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=0, line_spacing=2.0)

    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        set_run_font(run, bold=True, size=12)
    elif level == 2:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        set_run_font(run, bold=True, size=12)
    elif level == 3:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        set_run_font(run, bold=True, italic=True, size=12)

    return p


def add_body_paragraph(doc, text, indent_first_line=True):
    """Add a body paragraph with optional first-line indent."""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=0, line_spacing=2.0)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    if indent_first_line:
        p.paragraph_format.first_line_indent = Inches(0.5)

    # Parse markdown inline formatting: *italic* and **bold**
    add_formatted_runs(p, text)
    return p


def add_formatted_runs(paragraph, text):
    """Parse markdown inline formatting and add runs.
    Handles **bold**, *italic*, and plain text.
    """
    # Pattern: **bold** or *italic* or plain text
    pattern = r'(\*\*(.+?)\*\*|\*(.+?)\*|([^*]+))'
    parts = re.finditer(pattern, text)

    for match in parts:
        if match.group(2):  # **bold**
            run = paragraph.add_run(match.group(2))
            set_run_font(run, bold=True)
        elif match.group(3):  # *italic*
            run = paragraph.add_run(match.group(3))
            set_run_font(run, italic=True)
        elif match.group(4):  # plain text
            run = paragraph.add_run(match.group(4))
            set_run_font(run)


def add_block_quote(doc, text):
    """Add a block quote (indented paragraph for RQ/hypothesis statements)."""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=0, line_spacing=2.0)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(0)

    add_formatted_runs(p, text)
    return p


def add_reference(doc, text):
    """Add a reference entry with hanging indent."""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=0, line_spacing=2.0)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.5)

    # Parse for italic journal names and titles
    add_reference_formatted_runs(p, text)
    return p


def add_reference_formatted_runs(paragraph, text):
    """Parse reference text and italicize journal names / book titles.
    References use italic for content after the date until the next period
    for journal articles, or for book titles.
    We use a simpler approach: italicize content within *markers* if present,
    otherwise add as plain text.
    """
    # Check if there are any italic markers
    if "*" in text:
        add_formatted_runs(paragraph, text)
    else:
        run = paragraph.add_run(text)
        set_run_font(run)


def add_page_break(doc):
    """Add a page break."""
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(docx.oxml.OxmlElement("w:br"))
    from docx.oxml.ns import qn as qn2
    br = run._element.makeelement(qn2("w:br"), {qn2("w:type"): "page"})
    run._element.append(br)


def insert_page_break(doc):
    """Insert a proper page break."""
    p = doc.add_paragraph()
    p_elem = p._element
    pPr = p_elem.get_or_add_pPr()
    # Remove any existing pageBreakBefore
    for existing in pPr.findall(qn("w:pageBreakBefore")):
        pPr.remove(existing)
    pbr = p_elem.makeelement(qn("w:r"), {})
    br = p_elem.makeelement(qn("w:br"), {qn("w:type"): "page"})
    pbr.append(br)
    p_elem.append(pbr)
    return p


# ---------------------------------------------------------------------------
# Table and Figure insertion helpers
# ---------------------------------------------------------------------------

TABLE_TITLES = {
    1: "Table 1\nSample Characteristics by Generational Cohort (CPS-CEV 2017–2023, N = 201,168)",
    2: "Table 2\nSurvey-Weighted Logistic Regression Predicting Volunteering (Model 1: Socialization × Generation Interaction)",
    3: "Table 3\nFirst Step Effect: Average Marginal Effects for the Transition From No Socialization to Minimal Socialization, by Generation",
    4: "Table 4\nLatent Profile Analysis Model Fit Indices (EII Parameterization, 20,000 Subsample)",
    5: "Table 5\nLatent Profile Characteristics: Indicator Means and Volunteering Rates (6-Profile Solution, N = 197,497)",
    6: "Table 6\nVolunteering and Minimal Socialization Rates by Generation and Survey Wave (Survey-Weighted Estimates)",
}

TABLE_FILES = {
    1: "tables/table1_sample_characteristics.csv",
    2: "tables/table2_regression.csv",
    3: "tables/table3_first_step_ame.csv",
    4: "tables/table4_lpa_fit.csv",
    5: "tables/table5_lpa_profiles.csv",
    6: "tables/table_gen_wave.csv",
}

FIGURE_MAP = {
    1: ("figures/conceptual_framework.png",
        "Figure 1. Conceptual framework. The primary relationship between in-person socialization (IV) and volunteering (DV) is moderated by generational cohort. Three-way interactions test whether education, employment, civic social media use, and the COVID-19 period further condition this relationship. Three analytic stages provide methodological triangulation."),
    2: ("figures/fig1_pred_prob.png",
        "Figure 2. Predicted probability of volunteering by in-person socialization frequency and generation, from survey-weighted logistic regression (Model 1). The steep initial rise illustrates the First Step Effect; Gen Z's flattening beyond moderate socialization illustrates the generational plateau. Shaded bands represent 95% confidence intervals."),
    3: ("figures/fig4_ame_first_step.png",
        "Figure 3. The First Step Effect: Average Marginal Effect of transitioning from no socialization to minimal socialization (few times per year) on volunteering probability, by generation. The effect is universal across all five cohorts (7.3 to 10.3 percentage points). Error bars represent 95% confidence intervals."),
    4: ("figures/fig2_lpa_heatmap.png",
        "Figure 4. Civic engagement profile characteristics from latent profile analysis. All values are normalized to a 0\u2013100 scale for cross-indicator comparability. The green column shows volunteering rate (outcome variable, not used as an LPA indicator). Color intensity reflects within-indicator relative magnitude."),
    5: ("figures/fig3_lpa_gen_dist.png",
        "Figure 5. Generational distribution across six latent civic engagement profiles. Over half of Gen Z (50.6%) occupies the two socially disconnected profiles (Isolated Disengaged and Politically Aware Isolated), while only 3.7% reaches Fully Engaged status."),
}


TABLE_NOTES = {
    1: "Note. Survey-weighted estimates. Minimal socialization = reported socializing not at all or a few times per year. Civic SM = used social media for civic or political purposes. Test statistics are design-corrected F-tests.\n*p < .05. **p < .01. ***p < .001.",
    2: "Note. OR = odds ratio. CI = confidence interval. Survey-weighted logistic regression (quasibinomial). All models control for age, sex, race/ethnicity, education (BA+), employment, marital status, family income (log), metropolitan status, region, and post-COVID period. Interaction Wald tests use design-corrected F-statistics. McFadden pseudo R² computed as 1 − (deviance / null deviance).\n*p < .05. **p < .01. ***p < .001.",
    3: "Note. AME = average marginal effect; pp = percentage points; SE = standard error; CI = confidence interval. AMEs represent the estimated change in volunteering probability for the transition from \"Not at all\" to \"Few times/year\" socialization, computed via the marginaleffects package with survey weights.",
    4: "Note. EII = spherical, equal volume parameterization. BIC = Bayesian information criterion. ΔBIC = improvement from the previous model. Entropy reported for the selected model only. Model selection conducted on a random subsample (n = 20,000) and refit to the full sample (N = 197,497).",
    5: "Note. Profile labels assigned post hoc based on indicator patterns. Continuous indicators reported as M (SD); binary indicators as percentages. Volunteering rate = proportion who volunteered in the past 12 months (not used as an LPA indicator).",
    6: "Note. Survey-weighted estimates. Vol. = volunteering rate (%). No Soc. = reported socializing not at all (%). Low Soc. = reported socializing a few times per year or less (%). SE in parentheses. The 2021 wave was administered during the COVID-19 pandemic.",
}


def _set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    """Set individual cell borders. Each border param is a dict: {'sz': '4', 'val': 'single', 'color': '000000'} or None to remove."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = tc.makeelement(qn("w:tcBorders"), {})
        tcPr.append(tcBorders)
    for edge, attrs in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        # Remove existing
        for existing in tcBorders.findall(qn(f"w:{edge}")):
            tcBorders.remove(existing)
        if attrs:
            el = tc.makeelement(qn(f"w:{edge}"), {
                qn("w:val"): attrs.get("val", "single"),
                qn("w:sz"): attrs.get("sz", "4"),
                qn("w:space"): "0",
                qn("w:color"): attrs.get("color", "000000"),
            })
            tcBorders.append(el)
        else:
            # Set to none (remove border)
            el = tc.makeelement(qn(f"w:{edge}"), {
                qn("w:val"): "none",
                qn("w:sz"): "0",
                qn("w:space"): "0",
                qn("w:color"): "000000",
            })
            tcBorders.append(el)


def add_table_from_csv(doc, table_num):
    """Insert an APA 7th edition formatted table from CSV data.

    APA 7th table rules:
    - NO vertical lines anywhere
    - Horizontal lines: top of table, below header row, bottom of table ONLY
    - Table number: bold, italic (e.g., "Table 1")
    - Table title: italic, title case
    - Table note: flush left, "Note." in italic
    - Font: Times New Roman 10pt in cells, 12pt for title
    - Single-spaced within cells
    """
    csv_path = BASE_DIR / TABLE_FILES[table_num]
    if not csv_path.exists():
        p = doc.add_paragraph()
        run = p.add_run(f"[Table {table_num} ,  data file not found: {csv_path}]")
        set_run_font(run, italic=True)
        return

    # --- Table number (APA: bold, italic, on its own line) ---
    title_text = TABLE_TITLES.get(table_num, f"Table {table_num}")
    title_lines = title_text.split("\n")
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=12, after=0, line_spacing=2.0)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(title_lines[0])
    set_run_font(run, bold=True, italic=True, size=12)

    # --- Table title (APA: italic, title case, on next line) ---
    if len(title_lines) > 1:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=0, after=4, line_spacing=2.0)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(title_lines[1])
        set_run_font(run, italic=True, size=12)

    # --- Read CSV ---
    with open(csv_path, "r", encoding="utf-8") as f:
        reader = csv.reader(f)
        rows = list(reader)

    if not rows:
        return

    headers = rows[0]
    data_rows = rows[1:]
    n_cols = len(headers)
    n_rows = len(data_rows) + 1  # +1 for header

    # --- Create table (no built-in style  - we control all borders) ---
    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.style = "Normal Table"

    # Remove all default table borders at the table level
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else tbl._add_tblPr()

    # Set table width to 100% of page
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = tbl.makeelement(qn("w:tblW"), {})
        tblPr.append(tblW)
    tblW.set(qn("w:w"), "5000")
    tblW.set(qn("w:type"), "pct")

    # Remove table-level borders entirely
    tblBorders = tblPr.find(qn("w:tblBorders"))
    if tblBorders is not None:
        tblPr.remove(tblBorders)
    tblBorders = tbl.makeelement(qn("w:tblBorders"), {})
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        el = tbl.makeelement(qn(f"w:{edge}"), {
            qn("w:val"): "none", qn("w:sz"): "0",
            qn("w:space"): "0", qn("w:color"): "000000"
        })
        tblBorders.append(el)
    tblPr.append(tblBorders)

    # Define border styles
    RULE = {"val": "single", "sz": "8", "color": "000000"}  # thin black line
    NONE = None  # no border

    # --- Header row ---
    for j, header in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j > 0 else WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(header)
        set_run_font(run, bold=True, size=10)
        set_paragraph_spacing(p, before=2, after=2, line_spacing=1.0)
        # APA borders: top rule + bottom rule on header
        _set_cell_borders(cell, top=RULE, bottom=RULE, left=NONE, right=NONE)

    # --- Data rows ---
    for i, row_data in enumerate(data_rows):
        is_last_row = (i == len(data_rows) - 1)
        for j, val in enumerate(row_data):
            if j >= n_cols:
                continue
            cell = table.rows[i + 1].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(val)
            # Indent sub-rows (starting with spaces)
            is_subrow = val.startswith("  ")
            set_run_font(run, size=10)
            set_paragraph_spacing(p, before=1, after=1, line_spacing=1.0)
            # APA borders: only bottom rule on last row; no vertical lines ever
            _set_cell_borders(
                cell,
                top=NONE,
                bottom=RULE if is_last_row else NONE,
                left=NONE,
                right=NONE
            )

    # --- Table note (APA: flush left, "Note." italic, then regular text) ---
    note_text = TABLE_NOTES.get(table_num, "")
    if note_text:
        for note_line in note_text.split("\n"):
            p = doc.add_paragraph()
            set_paragraph_spacing(p, before=2, after=0, line_spacing=2.0)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            if note_line.startswith("Note."):
                run = p.add_run("Note. ")
                set_run_font(run, italic=True, size=10)
                run2 = p.add_run(note_line[6:])
                set_run_font(run2, size=10)
            elif note_line.startswith("*"):
                run = p.add_run(note_line)
                set_run_font(run, size=10)
            else:
                run = p.add_run(note_line)
                set_run_font(run, size=10)

    # Spacing after table
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=4, after=12, line_spacing=2.0)


def add_figure_inline(doc, figure_num):
    """Insert a figure with APA caption below."""
    if figure_num not in FIGURE_MAP:
        p = doc.add_paragraph()
        run = p.add_run(f"[Figure {figure_num} ,  not mapped]")
        set_run_font(run, italic=True)
        return

    img_path_rel, caption = FIGURE_MAP[figure_num]
    img_path = BASE_DIR / img_path_rel

    if not img_path.exists():
        p = doc.add_paragraph()
        run = p.add_run(f"[Figure {figure_num} ,  file not found: {img_path_rel}]")
        set_run_font(run, italic=True)
        return

    # Add some space before figure
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=6, after=0, line_spacing=1.0)

    # Insert image centered
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=0, after=4, line_spacing=1.0)
    run = p.add_run()
    run.add_picture(str(img_path), width=Inches(5.5))

    # Caption (APA: italic "Figure X." then regular text)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(p, before=4, after=12, line_spacing=2.0)
    # Split caption: "Figure X." is italic, rest is regular
    parts = caption.split(". ", 1)
    run = p.add_run(parts[0] + ". ")
    set_run_font(run, italic=True, size=12)
    if len(parts) > 1:
        run = p.add_run(parts[1])
        set_run_font(run, size=12)


# ---------------------------------------------------------------------------
# Document setup
# ---------------------------------------------------------------------------

def setup_document():
    """Create and configure the document with APA margins and defaults."""
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)

    # Set margins (1 inch all around for APA)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    return doc


# ---------------------------------------------------------------------------
# Section builders  - each adds heading then content immediately
# ---------------------------------------------------------------------------

def build_title_page(doc):
    """Build APA title page."""
    # Add some spacing at top
    for _ in range(3):
        p = doc.add_paragraph()
        set_paragraph_spacing(p, line_spacing=2.0)

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, line_spacing=2.0)
    run = p.add_run("Bowling Alone, Scrolling Together:")
    set_run_font(run, bold=True, size=12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, line_spacing=2.0)
    run = p.add_run("In-Person Socialization, Social Isolation, and Volunteering")
    set_run_font(run, bold=True, size=12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, line_spacing=2.0)
    run = p.add_run("Across Five Generational Cohorts")
    set_run_font(run, bold=True, size=12)

    # Blank line
    p = doc.add_paragraph()
    set_paragraph_spacing(p, line_spacing=2.0)

    # Authors
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, line_spacing=2.0)
    run = p.add_run("Hosung You & Suzanna R. Windon")
    set_run_font(run, size=12)

    # Affiliation
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, line_spacing=2.0)
    run = p.add_run("The Pennsylvania State University")
    set_run_font(run, size=12)

    # Page break
    insert_page_break(doc)


def build_abstract(doc):
    """Build the Abstract section from 00_abstract.md."""
    paper_dir = os.path.dirname(os.path.abspath(__file__))
    md_path = os.path.join(paper_dir, "00_abstract.md")

    with open(md_path, "r", encoding="utf-8") as f:
        content = f.read()

    # Parse: heading, body, keywords
    lines = content.strip().split("\n")
    body_lines = []
    keywords_text = ""
    for line in lines:
        stripped = line.strip()
        if stripped.startswith("# "):
            continue  # skip heading, we add it manually
        if stripped.startswith("**Keywords:**"):
            keywords_text = stripped.replace("**Keywords:**", "").strip()
            continue
        if stripped:
            body_lines.append(stripped)

    add_heading_apa(doc, "Abstract", level=1)
    abstract_text = " ".join(body_lines)
    add_body_paragraph(doc, abstract_text, indent_first_line=False)

    if keywords_text:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, line_spacing=2.0)
        p.paragraph_format.first_line_indent = Inches(0.5)
        run = p.add_run("Keywords: ")
        set_run_font(run, italic=True)
        run2 = p.add_run(keywords_text)
        set_run_font(run2)

    insert_page_break(doc)


def _build_introduction_hardcoded(doc):
    """DEPRECATED: Build the Introduction section with hardcoded text."""
    add_heading_apa(doc, "Introduction", level=1)

    paragraphs = [
        'In May 2023, the U.S. Surgeon General declared loneliness and social isolation a public health epidemic, documenting extensive consequences for mental health, physical well-being, and mortality (Office of the Surgeon General, 2023). Yet while the health effects of social disconnection have received considerable empirical attention, one domain remains underexamined: the civic consequences of the loneliness epidemic, and particularly its effects on volunteering. This omission is consequential. A quarter century after Putnam (2000) documented the broad decline of associational life in *Bowling Alone*, emerging evidence suggests that the erosion of social connection has continued, and accelerated among younger cohorts, with potential ramifications for the civic infrastructure that depends on volunteerism.',

        'Twenge et al. (2019) documented that in-person social interaction among U.S. adolescents declined substantially between 2003 and 2017, a trend that predates the COVID-19 pandemic and thus cannot be attributed solely to lockdown-era disruptions. If social interaction is a precondition for civic recruitment, as both social capital theory (Putnam, 2000) and the civic voluntarism model (Verba et al., 1995) suggest, then generational declines in face-to-face socialization may represent a structural barrier to volunteering that existing interventions fail to address.',

        'Understanding the civic consequences of social isolation is particularly urgent for the volunteering sector. Unlike voting or charitable giving, organized volunteering demands sustained investment of time and typically requires coordination with others, features that make it especially dependent on social networks and in-person contact. Yet the empirical literature on volunteering has treated social isolation primarily as a control variable rather than a central mechanism, and studies that do examine socialization focus on its general association with civic engagement without testing whether the relationship varies systematically across generational cohorts that experienced fundamentally different socialization environments.',

        'Despite growing attention to the health consequences of social isolation (Cacioppo & Cacioppo, 2014; Holt-Lunstad et al., 2010), its implications for civic participation, and particularly volunteering, remain empirically underexplored. Existing research on generational differences in volunteering tends to rely on simple cohort comparisons without examining how the *mechanisms* driving volunteering may differ across generations. This leaves a critical question unanswered: Is the decline in youth volunteering driven by changing values and attitudes, or by the structural erosion of the social connections that make volunteering possible in the first place?',

        'The present study addresses this gap using four waves of nationally representative CPS-CEV data (2017, 2019, 2021, 2023; *N* \u2248 201,000) to examine the association between in-person socialization frequency and volunteering across five generational cohorts. We employ an integrated analytic design comprising three complementary approaches, survey-weighted logistic regression, Latent Profile Analysis, and Gradient Boosting with TreeSHAP, to address three research questions:',
    ]

    for para_text in paragraphs:
        add_body_paragraph(doc, para_text)

    # RQ block quotes
    rq1 = '**RQ1** (Variable-Centered): How does the association between in-person socialization frequency and volunteering differ across generational cohorts, and does this relationship exhibit nonlinear threshold effects? Do education, employment, civic social media use, and the COVID-19 pandemic moderate this association differently by generation?'
    add_block_quote(doc, rq1)

    rq2 = '**RQ2** (Person-Centered): What distinct civic engagement typologies emerge from patterns of boycotting, contacting officials, political conversation, in-person socialization, organizational membership, and charitable donation, and how are generational cohorts distributed across these typologies?'
    add_block_quote(doc, rq2)

    rq3 = '**RQ3** (Predictive Validation): Does machine learning-based variable importance confirm socialization as the dominant predictor of volunteering, and do nonlinear threshold dynamics differ by generation?'
    add_block_quote(doc, rq3)

    # Final paragraph
    final = 'By spanning four survey waves that bracket the COVID-19 pandemic, this study can distinguish structural cohort effects from pandemic-specific disruptions, a critical distinction that single-wave analyses cannot make. By integrating variable-centered regression, person-centered profiling, and nonparametric machine learning, this study makes three contributions. First, it extends the Surgeon General\u2019s health-focused framing to a civic engagement context, demonstrating that social isolation is not merely a health risk but a barrier to civic participation. Second, it identifies a \u201cFirst Step Effect\u201d in which even minimal social contact produces the largest marginal gain in volunteering probability, a finding with direct implications for intervention design. Third, it reveals that conventional moderators of volunteering (education, social media use) operate differently for Generation Z than for older cohorts, suggesting that recruitment strategies calibrated to previous generations may be ineffective for the most socially isolated cohort.'
    add_body_paragraph(doc, final)


def _build_literature_review_hardcoded(doc):
    """DEPRECATED: Build the Literature Review with hardcoded text."""
    add_heading_apa(doc, "Literature Review", level=1)

    # --- Social Capital and Civic Recruitment (H2) ---
    add_heading_apa(doc, "Social Capital and Civic Recruitment", level=2)

    add_body_paragraph(doc,
        'The theoretical foundation for the relationship between social interaction and volunteering rests on social capital theory and the civic voluntarism model. Putnam (2000) defined social capital as the networks, norms, and trust that facilitate coordination and cooperation among individuals. In his framework, face-to-face social interaction is not merely a correlate of civic engagement but a generative mechanism: repeated contact builds the trust and reciprocity norms that lower the costs of collective action. Granovetter\u2019s (1973) weak ties theory extends this insight by arguing that peripheral social contacts, acquaintances encountered through occasional socialization, serve as bridges to new information and opportunities, including awareness of volunteer organizations and direct recruitment appeals.'
    )
    add_body_paragraph(doc,
        'The civic voluntarism model (CVM) identifies three categories of factors predicting civic participation: resources (time, money, civic skills), engagement (political interest, civic duty), and recruitment networks (Brady et al., 1995; Verba et al., 1995). The recruitment pathway is particularly relevant: individuals who socialize more frequently are more likely to encounter direct requests to volunteer, which the CVM identifies as the most proximate predictor of participation. Coleman (1988) further emphasized that social capital is embedded in the structure of relationships: closed social networks with dense interconnections facilitate the emergence of norms and sanctions that encourage prosocial behavior, including volunteering.'
    )

    # --- Generational Shifts in Socialization and Civic Participation (H2) ---
    add_heading_apa(doc, "Generational Shifts in Socialization and Civic Participation", level=2)

    add_body_paragraph(doc,
        'The theoretical expectation that social interaction facilitates civic engagement must be situated within evidence of dramatic generational shifts in socialization patterns. Twenge et al. (2019) documented that in-person social interaction among U.S. adolescents declined by approximately 30% between 2003 and 2017, a trend driven primarily by the displacement of face-to-face contact by digital communication. Twenge (2017) termed the post-1995 birth cohort \u201ciGen\u201d and documented widespread increases in loneliness, social isolation, and mental health challenges that emerged before, not as a consequence of, the COVID-19 pandemic.'
    )
    add_body_paragraph(doc,
        'These shifts intersect with broader transformations in civic engagement documented by scholars of political participation. Zukin et al. (2006) identified generational differences in civic participation, with younger Americans more likely to engage in individualized political consumption (boycotting, \u201cbuycotting\u201d) and less likely to participate in traditional institutional forms such as voting and formal volunteerism. Dalton (2008) proposed that this represents not disengagement but a shift from \u201cduty-based\u201d to \u201cengaged\u201d citizenship, where younger cohorts pursue civic goals through informal, often digital, channels. However, Flanagan and Levine (2010) and Wray-Lake and Hart (2012) documented that these alternative forms of engagement have not compensated for declines in traditional civic participation, particularly volunteering, among younger cohorts.'
    )
    add_body_paragraph(doc,
        'These theoretical perspectives generate competing predictions about the relationship between social isolation and volunteering across generations. If social capital theory is correct that face-to-face interaction is the primary mechanism through which civic recruitment occurs, then the dramatic decline in in-person socialization among Gen Z should produce a correspondingly large decline in volunteering, even after controlling for the traditional predictors emphasized by the CVM (education, income, civic skills). Alternatively, if digital communication has functionally replaced face-to-face contact as a recruitment channel, then the socialization\u2013volunteering link should be weaker for digitally native cohorts.'
    )

    # --- Why Generation Z? (H3) ---
    add_heading_apa(doc, "Why Generation Z?", level=3)

    add_body_paragraph(doc,
        'Generation Z is the focal cohort in this study for three reasons. First, Gen Z is the first generation to have entered adulthood entirely within a digitally mediated social environment, making them a critical test case for theories about whether online interaction can substitute for face-to-face socialization in generating civic participation. Second, descriptive data from the CPS-CEV consistently shows that Gen Z reports the highest rates of minimal socialization (CESOCIALIZE = \u201cNot at all\u201d), suggesting that if social isolation affects volunteering, the effects should be most visible in this cohort. Third, Gen Z represents the future volunteer workforce; understanding the mechanisms that facilitate or inhibit their civic participation has direct implications for nonprofit management, volunteer recruitment, and community development practice.'
    )
    add_body_paragraph(doc,
        'Generation is operationalized as a moderating variable rather than merely a demographic control because generational membership captures shared formative experiences, including exposure to digital technology, economic conditions during the transition to adulthood, and the COVID-19 pandemic, that shape how individuals relate to civic institutions and social networks (Hooghe, 2012; Kelle et al., 2025).'
    )

    # --- Education, Employment, and Civic Social Media (H2) ---
    add_heading_apa(doc, "Education, Employment, and Civic Social Media", level=2)

    add_body_paragraph(doc,
        'The civic voluntarism model identifies education and employment as primary resource predictors of civic participation (Verba et al., 1995). Higher education develops civic skills, expands social networks, and cultivates norms of civic duty (Schlozman et al., 2012). Employment provides organizational skills, access to workplace-based recruitment, and the financial resources that reduce the opportunity cost of volunteering (Musick & Wilson, 2008; Wilson, 2000). These resource pathways are well documented in the volunteering literature (Wilson, 2012).'
    )
    add_body_paragraph(doc,
        'However, whether these pathways operate with equal strength across generational cohorts has received limited attention. If Generation Z\u2019s social isolation reflects a structural shift in how young adults form social connections, then the conventional resource model, in which education and employment provide the skills and networks that facilitate volunteering, may be less effective for young adults who lack the foundational social contacts through which resources are converted into civic action.'
    )
    add_body_paragraph(doc,
        'We additionally examine civic use of social media as a potential moderating factor. Bennett and Segerberg (2012) theorized the emergence of \u201cconnective action\u201d, individualized, digitally mediated civic participation that bypasses traditional organizational structures. The \u201cdigital compensation\u201d hypothesis (Norris, 2002) suggests that online civic engagement could partially offset the volunteering deficit produced by social isolation, particularly for digitally native cohorts. If socially isolated individuals use social media for civic purposes (e.g., reading about community issues, sharing information about causes), this digital engagement might maintain awareness of volunteer opportunities and lower barriers to participation, even in the absence of face-to-face recruitment.'
    )

    # --- The Loneliness Epidemic and Civic Consequences (H2) ---
    add_heading_apa(doc, "The Loneliness Epidemic and Civic Consequences", level=2)

    add_body_paragraph(doc,
        'The Surgeon General\u2019s 2023 advisory situated social isolation within a public health framework, documenting associations between social disconnection and increased risk of heart disease, stroke, dementia, and premature mortality (Office of the Surgeon General, 2023). While the advisory focused primarily on health outcomes, it also noted that social isolation \u201cundermines the social infrastructure on which democratic participation depends\u201d, a claim that, to date, has received limited empirical testing in the volunteering context.'
    )
    add_body_paragraph(doc,
        'Cacioppo and Cacioppo (2014) provided a theoretical bridge between loneliness and civic withdrawal, arguing that perceived social isolation triggers a self-reinforcing cycle of hypervigilance to social threat, reduced trust, and withdrawal from social situations, precisely the conditions that inhibit civic recruitment. This theoretical account predicts that the relationship between social isolation and civic disengagement should be nonlinear: even small reductions in isolation (moving from complete social disconnection to occasional contact) may break the cycle of withdrawal and produce disproportionately large gains in civic participation. This prediction forms the basis for the \u201cFirst Step Effect\u201d hypothesis tested in this study.'
    )


def _build_purpose_rq_hardcoded(doc):
    """DEPRECATED: Build Purpose/RQ with hardcoded text."""
    add_heading_apa(doc, "Purpose and Research Questions", level=1)

    add_body_paragraph(doc,
        'The purpose of this study is to examine the association between in-person socialization frequency and volunteering across generational cohorts using four waves of nationally representative data from the Current Population Survey Civic Engagement and Volunteering Supplement (2017, 2019, 2021, 2023). To move beyond the limitations of any single analytic approach, we employ an integrated design that combines variable-centered regression, person-centered profiling, and machine learning validation. By spanning four waves that bracket the COVID-19 pandemic, this design allows us to distinguish structural cohort effects from pandemic-specific disruptions.'
    )
    add_body_paragraph(doc,
        'Figure 1 presents the conceptual model guiding this study. The primary relationship of interest is the association between in-person socialization frequency (independent variable) and volunteering status (dependent variable), moderated by generational cohort. Secondary moderation pathways examine whether education, employment status, civic social media use, and the COVID-19 pandemic condition the socialization\u2013volunteering relationship, and whether these moderating effects differ by generation (three-way interactions). The person-centered component identifies distinct typologies of civic engagement, while the machine learning component validates predictor importance without parametric assumptions.'
    )

    # RQ1
    add_body_paragraph(doc,
        '**Research Question 1** (Variable-Centered): How does the association between in-person socialization frequency and volunteering differ across generational cohorts, and do education, employment, civic social media use, and the COVID-19 pandemic moderate this association differently by generation?',
        indent_first_line=False
    )

    hypotheses_rq1 = [
        '*H1.* Higher levels of in-person socialization will be positively associated with volunteering across all generational cohorts.',
        '*H2.* Generation Z will exhibit a weaker socialization\u2013volunteering association at higher socialization levels compared with older cohorts (i.e., a plateau effect).',
        '*H3a.* Education will positively moderate the socialization\u2013volunteering association, but this moderating effect will be weaker for Generation Z than for older cohorts.',
        '*H3b.* Employment will positively moderate the socialization\u2013volunteering association, with effects that are relatively consistent across generations.',
        '*H3c.* Civic social media use will positively moderate the socialization\u2013volunteering association, but the moderating effect will be stronger for older cohorts than for Generation Z.',
        '*H3d.* The socialization\u2013volunteering association will be stable across the pre- and post-COVID periods, consistent with a cohort interpretation rather than a period effect.',
    ]
    for h in hypotheses_rq1:
        add_block_quote(doc, h)

    # RQ2
    add_body_paragraph(doc,
        '**Research Question 2** (Person-Centered): What distinct civic engagement typologies emerge from patterns of boycotting, contacting officials, political conversation, in-person socialization, organizational membership, and charitable donation, and how are generational cohorts distributed across these typologies?',
        indent_first_line=False
    )

    hypotheses_rq2 = [
        '*H4.* Generation Z will be disproportionately concentrated in low-engagement profiles characterized by minimal socialization and low organizational membership.',
        '*H5.* High-engagement profiles will show higher volunteering rates and be overrepresented among older cohorts.',
    ]
    for h in hypotheses_rq2:
        add_block_quote(doc, h)

    # RQ3
    add_body_paragraph(doc,
        '**Research Question 3** (Predictive Validation): Does machine learning-based variable importance confirm socialization as the dominant predictor of volunteering, and do nonlinear SHAP dependence patterns differ by generation?',
        indent_first_line=False
    )

    hypotheses_rq3 = [
        '*H6.* In-person socialization will be the single most important predictor of volunteering in a gradient boosting model, even when compared against education, income, and age.',
        '*H7.* SHAP dependence plots will reveal a steeper marginal effect at low socialization levels (1\u21922 transition) that flattens at higher levels, consistent with the \u201cFirst Step Effect.\u201d',
    ]
    for h in hypotheses_rq3:
        add_block_quote(doc, h)


def _build_method_hardcoded(doc):
    """DEPRECATED: Build Method with hardcoded text."""
    add_heading_apa(doc, "Method", level=1)

    # --- Data Source and Sample ---
    add_heading_apa(doc, "Data Source and Sample", level=2)

    add_body_paragraph(doc,
        'Data come from the Current Population Survey Civic Engagement and Volunteering Supplement (CPS-CEV), obtained via IPUMS-CPS (Flood et al., 2023). The CPS-CEV is administered as a supplement to the September CPS by the U.S. Census Bureau in collaboration with AmeriCorps. Approximately 57,000 households participate in each wave, making it the largest nationally representative survey measuring both volunteering behavior and civic engagement indicators simultaneously.'
    )
    add_body_paragraph(doc,
        'The CPS-CEV is uniquely suited to this study because it is the only nationally representative survey that simultaneously measures both volunteering behavior and in-person socialization frequency across a sample large enough to support subgroup analyses by generation. We pooled four waves of CPS-CEV data: September 2017, 2019, 2021, and 2023. These waves bracket the COVID-19 pandemic, allowing us to examine whether the socialization\u2013volunteering relationship changed across this major disruption to social life. The analytic sample includes all U.S. adults aged 18 and older with valid supplement weights (VLSUPPWT > 0) and nonmissing values on the dependent variable (VLSTATUS \u2208 {1, 2}), yielding *N* \u2248 201,000 spanning five generational cohorts.'
    )
    add_body_paragraph(doc,
        'Respondents with missing values on the key independent variable (CESOCIALIZE) were excluded from models involving socialization (less than 1% of the analytic sample). For moderating variables, cases with missing values were excluded listwise within each model. All analyses incorporate the supplement weight (VLSUPPWT) to produce nationally representative estimates.'
    )

    # --- Measures ---
    add_heading_apa(doc, "Measures", level=2)

    # Dependent Variable
    add_heading_apa(doc, "Dependent Variable", level=3)
    add_body_paragraph(doc,
        'Volunteering status was measured using VLSTATUS, which asks whether the respondent \u201cdid any volunteer activities through or for an organization\u201d in the past 12 months (1 = volunteered, 2 = did not volunteer). This measure was dichotomized for logistic regression and treated as an external validation variable for the latent profile analysis.'
    )

    # Key Independent Variable
    add_heading_apa(doc, "Key Independent Variable", level=3)
    add_body_paragraph(doc,
        'In-person socialization frequency was measured using CESOCIALIZE, which asks respondents how often they \u201cgot together socially with friends, relatives, or neighbors\u201d in the past 12 months. Response options range from 1 = \u201cNot at all\u201d to 6 = \u201cBasically every day.\u201d For logistic regression, socialization is entered as a factor variable with \u201cNot at all\u201d as the reference category, allowing estimation of nonlinear threshold effects at each transition. For the GBM analysis, it is entered as a continuous ordinal predictor.'
    )

    # Generational Cohorts
    add_heading_apa(doc, "Generational Cohorts", level=3)
    add_body_paragraph(doc,
        'Respondents were classified into five generational cohorts based on birth year (calculated as survey year minus age): Generation Z (born 1997 or later), Millennials (1981\u20131996), Generation X (1965\u20131980), Baby Boomers (1946\u20131964), and the Silent Generation (born before 1946). Generation Z serves as the reference category in regression models.'
    )

    # Moderating Variables
    add_heading_apa(doc, "Moderating Variables", level=3)
    add_body_paragraph(doc,
        'Four theoretically motivated moderators were examined. *Education* was measured using EDUC (IPUMS harmonized) and dichotomized as bachelor\u2019s degree or higher (EDUC \u2265 111) versus less than a bachelor\u2019s degree, consistent with the civic voluntarism model\u2019s emphasis on educational thresholds. *Employment status* was dichotomized as currently employed (EMPSTAT \u2208 {10, 12}) versus not employed. *Civic social media use* was measured using VLSOCMEDIA, which asks how often the respondent \u201cposted, shared, or discussed issues on civic or political topics on social media\u201d (1 = basically every day through 6 = not at all). This was dichotomized as any civic social media use (VLSOCMEDIA \u2264 4) versus no use, following the conceptual distinction between digital civic engagement and non-engagement.'
    )
    add_body_paragraph(doc,
        '*COVID-19 period* was operationalized as a binary variable distinguishing pre-pandemic waves (2017, 2019) from post-pandemic waves (2021, 2023). This coding captures the broad disruption to social life rather than a precise epidemiological boundary, and is included to test whether the socialization\u2013volunteering relationship is a stable structural feature or was altered by pandemic-related changes to social behavior.'
    )

    # Latent Profile Analysis Indicators
    add_heading_apa(doc, "Latent Profile Analysis Indicators", level=3)
    add_body_paragraph(doc,
        'Six indicators were used to construct civic engagement typologies: (1) *Boycotting* (CEBOYCOTT): whether the respondent boycotted a product for social or political reasons (binary); (2) *Contacting officials* (CEPUBOFF): whether the respondent contacted a public official (binary); (3) *Political conversation* (CEPOLCONV): frequency of discussing politics (ordinal, 1\u20136); (4) *In-person socialization* (CESOCIALIZE): frequency of in-person social contact (ordinal, 1\u20136); (5) *Organizational membership* (VLMEMBERN): number of organizations to which the respondent belongs (count; respondents reporting no membership via VLMEMBER were coded as 0); and (6) *Charitable donation* (VLDONATE): whether the respondent donated to a charitable or religious organization (binary). All six indicators are available across all four survey waves with approximately 94% valid response rates. This set captures multiple dimensions of civic life, political action, social connection, institutional membership, and philanthropic behavior, enabling person-centered identification of distinct engagement patterns.'
    )

    # Control Variables
    add_heading_apa(doc, "Control Variables", level=3)
    add_body_paragraph(doc,
        'Models controlled for age (continuous), sex (female = 1), race/ethnicity (White non-Hispanic [reference], Black non-Hispanic, Hispanic, Asian non-Hispanic, other), marital status (married = 1), family income (log-transformed midpoint), metropolitan status (metro = 1), and Census region (Northeast, Midwest, South, West [reference]). In Models 1\u20134, the post-COVID indicator also serves as a control; in Model 5, it is the focal moderator.'
    )

    # --- Analytic Strategy ---
    add_heading_apa(doc, "Analytic Strategy", level=2)
    add_body_paragraph(doc,
        'The analysis proceeds through three integrated stages, each addressing a distinct research question while providing methodological triangulation.'
    )

    # Stage 1
    add_heading_apa(doc, "Stage 1: Survey-Weighted Logistic Regression (RQ1)", level=3)
    add_body_paragraph(doc,
        'To address RQ1, we estimate survey-weighted logistic regression models predicting volunteering status. Socialization frequency is entered as a categorical predictor to capture nonlinear effects without imposing functional form assumptions. Model 1 includes the socialization \u00d7 generation interaction with post-COVID as a control. Models 2\u20134 extend to three-way interactions: socialization \u00d7 generation \u00d7 education (Model 2), socialization \u00d7 generation \u00d7 employment (Model 3), and socialization \u00d7 generation \u00d7 civic social media (Model 4). Model 5 tests the socialization \u00d7 generation \u00d7 post-COVID interaction to assess whether the socialization\u2013volunteering relationship changed across the pandemic period.'
    )
    add_body_paragraph(doc,
        'We report odds ratios and average marginal effects (AMEs) computed via the *marginaleffects* package (Arel-Bundock, 2023). AMEs represent the average change in the predicted probability of volunteering for a one-unit change in the predictor, averaged across the covariate distribution. Wald tests assess the statistical significance of interaction terms.'
    )

    # Stage 2
    add_heading_apa(doc, "Stage 2: Latent Profile Analysis (RQ2)", level=3)
    add_body_paragraph(doc,
        'To address RQ2, we employ Latent Profile Analysis (LPA) to identify distinct civic engagement typologies using six standardized indicators from the pooled four-wave sample. LPA is a model-based clustering technique that identifies subgroups of individuals with similar patterns across multiple indicators (Collins & Lanza, 2010; Muth\u00e9n & Muth\u00e9n, 2000). We evaluate models with 2\u20137 profiles using varying variances and zero covariances (Vermunt & Magidson, 2002), selecting the optimal number based on Bayesian Information Criterion (BIC), entropy, and Bootstrap Likelihood Ratio Test (BLRT). Given the large sample size (*N* > 200,000), initial model selection is conducted on a random subsample of 20,000 cases, with the selected model refit to the full sample.'
    )
    add_body_paragraph(doc,
        'After identifying profiles, we characterize each by its indicator means on the original scale, examine the generational distribution across profiles, compare volunteering rates by profile and generation, and assess whether profile distributions shifted between pre- and post-COVID periods.'
    )

    # Stage 3
    add_heading_apa(doc, "Stage 3: Gradient Boosting with TreeSHAP (RQ3)", level=3)
    add_body_paragraph(doc,
        'To address RQ3, we train an XGBoost gradient boosting classifier (Friedman, 2001) predicting volunteering from the same predictors used in the regression models, including post-COVID as a feature. This nonparametric approach makes no assumptions about functional form, automatically captures interactions, and provides a complementary test of predictor importance. Model performance is evaluated using AUC on a held-out test set (20% split).'
    )
    add_body_paragraph(doc,
        'Feature importance is decomposed using TreeSHAP (Lundberg & Lee, 2017; Lundberg et al., 2020), which assigns each predictor a Shapley value representing its marginal contribution to each prediction. We report mean |SHAP| values for the full sample and by generation, and generate dependence plots for socialization that reveal threshold effects. This validates whether the patterns observed in the parametric regression, particularly the First Step Effect and generational variation, are also captured by a nonlinear model without imposed interaction structure.'
    )


def build_section_from_md(doc, md_path):
    """Read a markdown file and build the section in the DOCX.
    Handles: # Heading 1, ## Heading 2, ### Heading 3, and body paragraphs.
    Skips HTML comments and blank lines. Numbered lists become body paragraphs.
    """
    import os
    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.readlines()

    paragraphs = []
    current_para = []

    for line in lines:
        stripped = line.strip()

        # Handle TABLE/FIGURE insertion markers
        table_match = re.match(r'<!--\s*TABLE:(\d+)\s*-->', stripped)
        figure_match = re.match(r'<!--\s*FIGURE:(\d+)\s*-->', stripped)
        if table_match:
            if current_para:
                paragraphs.append(("body", " ".join(current_para)))
                current_para = []
            paragraphs.append(("table", int(table_match.group(1))))
            continue
        if figure_match:
            if current_para:
                paragraphs.append(("body", " ".join(current_para)))
                current_para = []
            paragraphs.append(("figure", int(figure_match.group(1))))
            continue

        # Skip other HTML comments
        if stripped.startswith("<!--") and "-->" in stripped:
            continue

        # Blank line = paragraph break
        if not stripped:
            if current_para:
                paragraphs.append(("body", " ".join(current_para)))
                current_para = []
            continue

        # Headings
        if stripped.startswith("### "):
            if current_para:
                paragraphs.append(("body", " ".join(current_para)))
                current_para = []
            paragraphs.append(("h3", stripped[4:]))
        elif stripped.startswith("## "):
            if current_para:
                paragraphs.append(("body", " ".join(current_para)))
                current_para = []
            paragraphs.append(("h2", stripped[3:]))
        elif stripped.startswith("# "):
            if current_para:
                paragraphs.append(("body", " ".join(current_para)))
                current_para = []
            paragraphs.append(("h1", stripped[2:]))
        elif stripped.startswith(">"):
            if current_para:
                paragraphs.append(("body", " ".join(current_para)))
                current_para = []
            paragraphs.append(("quote", stripped.lstrip("> ")))
        else:
            # Strip numbered list markers (e.g., "1. **Boycotting**")
            import re as _re
            cleaned = _re.sub(r"^\d+\.\s+", "", stripped)
            current_para.append(cleaned)

    if current_para:
        paragraphs.append(("body", " ".join(current_para)))

    for kind, text in paragraphs:
        if kind == "h1":
            add_heading_apa(doc, text, level=1)
        elif kind == "h2":
            add_heading_apa(doc, text, level=2)
        elif kind == "h3":
            add_heading_apa(doc, text, level=3)
        elif kind == "quote":
            add_block_quote(doc, text)
        elif kind == "table":
            add_table_from_csv(doc, text)  # text is table number
        elif kind == "figure":
            add_figure_inline(doc, text)  # text is figure number
        elif kind == "body" and text.strip():
            add_body_paragraph(doc, text)


def _build_results_hardcoded(doc):
    """DEPRECATED: Build the Results section with hardcoded text."""
    add_heading_apa(doc, "Results", level=1)

    results_paragraphs = [
        'The analysis drew on 201,168 respondents pooled across four waves of the Current Population Survey Volunteer Supplement (2017, 2019, 2021, 2023), comprising 111,965 pre-COVID and 89,203 post-COVID observations. The overall volunteering rate was 32.0 percent. Generational subsamples ranged from 11,038 (Gen Z) to 68,065 (Baby Boomers), with Millennials (50,325) and Gen X (48,042) forming the two intermediate cohorts and the Silent Generation contributing 23,698 cases. Table 1 reports sample characteristics by generation. Demographic profiles conformed to expectations: Gen Z respondents were youngest (M = 21.1 years), least likely to hold a bachelor\u2019s degree (14.3%), and least likely to be married (9.6%), whereas Boomers and Silent Generation respondents exhibited higher rates of educational attainment relative to their cohort norms, marriage, and residential stability outside metropolitan areas. Female representation was roughly balanced across cohorts, ranging from 49.9 percent (Gen Z) to 57.1 percent (Silent).',

        'The descriptive volunteering gradient across generations followed an inverted-U pattern: Gen Z reported the second-lowest rate (23.9%), trailing only the Silent Generation (25.9%), while Gen X exhibited the highest rate (36.8%), followed by Boomers (32.2%) and Millennials (31.8%). More consequential for the present inquiry was the socialization gradient. Mean socialization scores increased monotonically from Gen Z (M = 3.30) to the Silent Generation (M = 4.36), but the critical indicator was the proportion reporting minimal socialization, socializing with neighbors, friends, or family not more than a few times per year (coded as CESOCIALIZE = 1). Nearly half of Gen Z respondents (46.3%) fell into this minimal-socialization category, a rate more than double that of Boomers (19.0%) and the Silent Generation (19.5%). Table 2 reveals that this pattern was remarkably stable across survey waves: Gen Z minimal socialization hovered between 44.6 and 47.8 percent from 2017 to 2023, showing no meaningful pre-to-post-COVID shift. In contrast, the Silent Generation exhibited a gradual upward drift from 17.6 percent in 2017 to 24.9 percent in 2023, consistent with age-related declines in social opportunity. Table 3 presents raw volunteering rates by socialization level and generation. Within every generation, volunteering increased with socialization frequency, but the gradient was notably shallower for Gen Z: the difference between the lowest and highest socialization categories was 17.6 percentage points for Gen Z compared with 25.4 points for Gen X and 26.2 points for Boomers. Table 4 documents post-COVID declines in volunteering across all generations, with Gen X experiencing the largest absolute decrease (37.5% to 31.5%, a 6.0 percentage-point drop) and Gen Z the smallest (26.6% to 21.8%, a 4.8-point drop). Minimal socialization rates increased modestly for older cohorts after COVID (e.g., Silent: 19.5% to 24.2%) but were essentially stable for Gen Z (46.4% to 47.8%), reinforcing the interpretation that Gen Z\u2019s social disconnection preceded and was structurally independent of pandemic disruption.',

        'To address the first research question, whether the relationship between face-to-face socialization and volunteering varies across generations, a series of survey-weighted logistic regression models was estimated. Model 1 regressed the binary volunteering indicator on the six-level socialization factor, generation, and their interaction, controlling for age, sex, education, family income, employment status, marital status, metropolitan residence, charitable giving, civic social media use, and survey wave. The socialization-by-generation interaction was statistically significant (Wald F = 5.54, df = 20, *p* < .001), confirming that the socialization\u2013volunteering relationship is not uniform across generational cohorts.',

        'Figure 2 displays the predicted probabilities of volunteering across socialization levels for each generation, derived from Model 1. The most striking pattern was the Gen Z plateau. Among Gen Z respondents, the predicted probability of volunteering rose sharply from 27.3 percent at the lowest socialization level (not at all) to 48.4 percent at moderate socialization (a few times per month), a 21.1 percentage-point increase, but then effectively flatlined: the predicted probability was 48.3 percent for those socializing a few times per week and 50.8 percent for those socializing daily. This ceiling effect meant that the highest-socializing Gen Z respondents achieved volunteering probabilities no greater than those of moderately socializing Gen Z respondents, a pattern absent in all other generations. By contrast, Boomers exhibited a more continuously rising trajectory, climbing from 16.7 percent at the lowest level to 46.8 percent at the highest, a 30.1 percentage-point range with no evidence of a ceiling. Gen X displayed the steepest overall gradient, moving from 22.0 percent to 48.1 percent, and showed sustained gains even at the highest socialization frequencies. The Silent Generation exhibited the lowest predicted probabilities throughout but maintained a consistently positive slope from 13.7 percent to 40.9 percent. Millennials occupied an intermediate position, rising from 21.7 percent to 45.6 percent with a mild deceleration at the upper end that was less pronounced than Gen Z\u2019s plateau.',

        'The Gen Z plateau warrants careful interpretation. It suggests that for the youngest cohort, the marginal return to volunteering from additional face-to-face socialization beyond a moderate threshold is negligible once covariates are held constant. One mechanism consistent with this pattern is that Gen Z individuals who socialize at moderate-to-high frequencies may already be embedded in social networks that are qualitatively different from those of older cohorts, networks that are more digitally mediated, more geographically dispersed, and less anchored in the community-based institutions (religious congregations, civic associations, neighborhood groups) that historically channel social contact into volunteer recruitment. Under this interpretation, additional in-person contact beyond a moderate level does not generate additional exposure to volunteering opportunities because the social contexts in which high-frequency Gen Z socialization occurs are decoupled from formal voluntarism.',

        'Average marginal effects (AMEs) were computed for each pairwise transition along the socialization scale, separately by generation. The most substantively important finding was the First Step Effect: the transition from no socialization to minimal socialization (not at all to a few times per year) produced the single largest marginal gain in volunteering probability for every generation. For Gen Z, this first step yielded an AME of 8.1 percentage points (95% CI [5.1, 11.1]); for Millennials, 8.0 points [6.5, 9.4]; for Gen X, 10.3 points [8.7, 11.9]; for Boomers, 8.8 points [7.4, 10.1]; and for the Silent Generation, 7.3 points [5.3, 9.4]. The magnitude of this first step was remarkably consistent across generations, with Gen X showing the largest gain and the Silent Generation the smallest. Subsequent transitions yielded progressively smaller marginal effects, consistent with a concave dose\u2013response relationship in which the initial move from social isolation to even minimal social contact confers the greatest benefit for civic participation. The universality of this first step effect across cohorts indicates that breaking out of complete social isolation is a generationally invariant mechanism for facilitating volunteering, even as the shape of the socialization\u2013volunteering curve at higher levels varies substantially by generation.',

        'Models 2 through 5 tested whether the socialization-by-generation interaction was further moderated by education, employment status, civic social media use, and the COVID period, respectively. Each model added a three-way interaction term to the base specification while retaining all lower-order terms and covariates.',

        'Model 2, which introduced the education-by-socialization-by-generation interaction, yielded a statistically significant three-way term (*p* = .005). The substantive pattern indicated that the socialization\u2013volunteering gradient was steeper among respondents holding at least a bachelor\u2019s degree, but this amplification was not uniform across generations. For Gen X and Boomers with bachelor\u2019s degrees, the predicted probability of volunteering at the highest socialization level exceeded 55 percent, whereas the corresponding figure for Gen Z degree holders was approximately 52 percent, still subject to the plateau effect observed in Model 1. Among respondents without a bachelor\u2019s degree, the generational differences were attenuated, with all cohorts showing relatively flat socialization\u2013volunteering gradients at higher socialization levels. This finding suggests that education functions as a resource multiplier for the socialization\u2013volunteering pathway, amplifying the returns to social contact, but that this amplification is partially constrained for Gen Z by the same ceiling dynamics documented above.',

        'Model 3 tested the employment-by-socialization-by-generation interaction, which reached marginal significance (*p* = .049). The pattern suggested that employment moderates the socialization\u2013volunteering relationship differently across generations, with employed Gen Z and Millennial respondents showing slightly steeper gradients than their non-employed counterparts. This is consistent with the workplace serving as a supplementary channel for volunteer recruitment that partially compensates for the weaker community-based recruitment pathways available to younger cohorts. However, the marginal significance of this interaction warrants caution in interpretation, and the effect sizes were modest.',

        'Model 4 introduced the civic social media-by-socialization-by-generation interaction, which was not statistically significant (*p* = .964). This null finding is important because it speaks directly to the digital compensation hypothesis. If civic social media use substituted for face-to-face socialization in facilitating volunteering, one would expect the three-way interaction to be significant, with the socialization\u2013volunteering gradient being flatter among civic social media users (because digital engagement would partially replace the role of in-person contact). The absence of any such moderation indicates that civic social media and face-to-face socialization operate through independent pathways: digital civic engagement does not buffer against the consequences of low in-person socialization for volunteering, nor does it enhance the returns to face-to-face contact. A supplementary descriptive analysis among socially isolated respondents (those reporting minimal socialization) provided additional texture. Among isolated Gen Z respondents, those who used social media for civic purposes volunteered at a rate of 18.1 percent compared to 15.0 percent among non-users, a modest 3.1 percentage-point difference. The corresponding differences were larger for older cohorts: 7.7 points for Millennials (23.6% vs. 15.9%), 8.0 points for Gen X (24.3% vs. 16.3%), 8.4 points for Boomers (20.2% vs. 11.8%), and 9.9 points for the Silent Generation (18.6% vs. 8.7%). Thus, while civic social media use was associated with somewhat higher volunteering rates among the socially isolated, this association was weakest for Gen Z, the generation most likely to be digitally native, further undermining the digital compensation hypothesis.',

        'Model 5 tested the COVID-by-socialization-by-generation interaction, which was not statistically significant (*p* = .344). This null result is among the most consequential findings of the study because it establishes that the generationally differentiated socialization\u2013volunteering relationship documented in Model 1 is temporally stable across the pre- and post-COVID periods. If the observed patterns were artifacts of pandemic-induced disruptions to social life, one would expect the three-way interaction to be significant, reflecting differential changes in the socialization\u2013volunteering gradient across generations from pre- to post-COVID periods. The absence of such an interaction supports a cohort-based interpretation: the Gen Z plateau and the generational variation in the socialization\u2013volunteering gradient reflect enduring differences in how generational cohorts convert social contact into civic participation, rather than transient responses to an exogenous shock. The pre-COVID AME for Gen Z\u2019s first step was 10.7 percentage points (95% CI [5.8, 15.6]), and the overall relationship structure was preserved in the post-COVID period, reinforcing the stability of this mechanism. Taken together, Models 1 through 5 establish that face-to-face socialization is a significant and generationally variable predictor of volunteering, that the relationship is characterized by a pronounced plateau for Gen Z, that education amplifies the gradient while employment provides marginal additional moderation, and that neither civic social media use nor the COVID pandemic meaningfully alters this generational pattern.',

        'Turning to the second research question, whether distinct profiles of civic engagement exist and whether generational membership predicts profile assignment, a latent profile analysis was estimated on six civic engagement indicators: boycotting, contacting public officials, political conversation frequency, socialization frequency, organizational membership count, and charitable giving. Model selection was guided by Bayesian Information Criterion, with the equal-volume spherical (EII) parameterization providing optimal fit. The BIC improvement from five to six classes was the largest single-step improvement observed (\u0394BIC = 13,977), motivating the selection of the six-class solution. Classification quality was strong, with an entropy value of 0.893 and 93.0 percent of cases assigned to their modal class with a posterior probability of 0.70 or greater.',

        'The six profiles, ordered by overall engagement intensity, were labeled as follows (Table 5; Figure 3). Profile 1, Isolated Disengaged (19.0%, *N* = 37,568), was characterized by minimal scores on all indicators: members of this class reported almost no boycotting (1.4%), negligible contact with public officials (0.7%), very low political conversation (M = 1.31), very low socialization (M = 1.34), near-zero organizational membership (M = 0.08), low charitable giving (15.0%), and the lowest volunteering rate in the sample (9.0%). This profile represents a segment of the population that is essentially disconnected from all forms of civic life. Profile 2, Politically Aware Isolated (7.1%, *N* = 14,089), shared Profile 1\u2019s social isolation (socialization M = 1.51) and low organizational involvement (M = 0.15) but was distinguished by elevated political conversation (M = 4.91) and modest volunteering (11.7%). Zero percent of this class reported boycotting, and none reported charitable giving, suggesting a profile of individuals who are cognitively engaged with politics but socially and behaviorally disengaged from community-level participation. Profile 3, Socially Active Non-Donors (17.5%, *N* = 34,520), exhibited the highest socialization scores (M = 4.70) of any class but combined this social activity with the absence of charitable giving (0%) and low volunteering (14.6%). This profile challenges the assumption that social connectedness automatically translates into civic participation, suggesting that high socialization without institutional anchoring may not channel social capital into formal voluntarism. Profile 4, Mainstream Donors (33.1%, *N* = 65,417), was the largest class and represented a conventional civic engagement pattern: moderate socialization (M = 3.77), moderate political conversation (M = 3.60), near-universal charitable giving (99.9%), and a volunteering rate of 41.4 percent. This class also reported meaningful organizational membership (M = 0.71) and no boycotting or public official contact, consistent with an institutionally embedded but politically quiescent civic style. Profile 5, Activist Boycotters (11.3%, *N* = 22,266), was defined by universal boycotting (100%), elevated political conversation (M = 4.28), substantial organizational membership (M = 1.03), and a high volunteering rate (46.9%). Profile 6, Fully Engaged (12.0%, *N* = 23,637), scored highest on virtually every indicator: universal contact with public officials (100%), frequent boycotting (49.7%), the highest political conversation scores (M = 4.59), high socialization (M = 4.04), the most extensive organizational memberships (M = 1.81), robust charitable giving (86.6%), and the highest volunteering rate (66.7%).',

        'The generational distribution across these profiles was strongly patterned (Table 6; Figure 4). Gen Z was dramatically overrepresented in the Isolated Disengaged profile: 35.3 percent of Gen Z respondents were assigned to Profile 1, compared with 22.4 percent of Millennials, 18.1 percent of Gen X, 15.3 percent of Boomers, and 16.8 percent of the Silent Generation. Conversely, Gen Z was steeply underrepresented in the Fully Engaged profile, with only 3.7 percent classified there versus 15.1 percent of Boomers and 12.5 percent of Gen X. An additional 15.3 percent of Gen Z fell into Profile 2 (Politically Aware Isolated), meaning that over half of Gen Z respondents (50.6%) occupied one of the two profiles characterized by social isolation. The Mainstream Donor pathway (Profile 4) attracted only 16.4 percent of Gen Z compared with 41.9 percent of the Silent Generation and 35.9 percent of Boomers, indicating that Gen Z has not yet entered, or may not enter, the institutionally anchored civic engagement pattern that has historically been the modal form of American voluntarism. Gen Z\u2019s representation in the Activist Boycotter profile (8.9%) was lower than that of Millennials (12.4%) or Gen X (12.3%), suggesting that politically expressive civic engagement is not yet a dominant pathway for the youngest cohort.',

        'The pre- and post-COVID profile distributions for Gen Z provided further evidence of structural stability. The proportion of Gen Z in the Isolated Disengaged profile was essentially unchanged from pre-COVID (35.5%) to post-COVID (35.2%), as was their representation in the Fully Engaged profile (3.7% in both periods). The most notable shift was a doubling of Gen Z\u2019s presence in the Activist Boycotter profile from 5.9 percent pre-COVID to 10.7 percent post-COVID, accompanied by a modest decline in the Socially Active Non-Donors profile (22.3% to 19.4%) and the Politically Aware Isolated profile (16.7% to 14.3%). These patterns suggest that while the pandemic may have catalyzed a shift toward political activism among a subset of Gen Z, it did not materially alter the fundamental finding that more than a third of this generation occupies a profile of comprehensive civic disengagement. The stability of the Isolated Disengaged proportion across the COVID divide, echoing the temporal stability documented in the regression analyses, provides convergent evidence that Gen Z\u2019s social disconnection is a cohort characteristic rather than a pandemic artifact.',

        'The third research question asked whether the relative importance of face-to-face socialization as a predictor of volunteering differs across generations. A gradient boosting machine was trained on the full covariate set using a generation-stratified approach, with separate SHAP (SHapley Additive exPlanations) value decompositions computed for each generational subsample. The overall model achieved a test-set area under the receiver operating characteristic curve (AUC) of 0.731, indicating acceptable discriminative performance for a binary outcome with a 32 percent base rate and a diverse set of predictors.',

        'Table 7 and Figure 5 present the mean absolute SHAP values for each feature by generation. The central finding was that socialization frequency (CESOCIALIZE) ranked as the single most important predictor of volunteering for Gen Z (mean |SHAP| = 0.467) but ranked second for every other generation. For Millennials, Gen X, Boomers, and the Silent Generation, education was the most important predictor (mean |SHAP| values of 0.426, 0.472, 0.449, and 0.464, respectively), with socialization consistently occupying the second position (0.391, 0.398, 0.414, 0.447). The reversal in feature importance rankings for Gen Z is notable because it held even after accounting for age, which was the third-ranked predictor for Gen Z (mean |SHAP| = 0.246) and thus could not explain away the primacy of socialization. Family income, gender, civic social media use, and the COVID indicator contributed modestly and relatively uniformly across generations, with mean absolute SHAP values ranging from 0.117 to 0.186.',

        'SHAP dependence plots (Figure 6) provided additional granularity on the functional form of the socialization\u2013volunteering relationship. Consistent with the regression-based AME analysis, the steepest SHAP gradient occurred at the transition from the lowest socialization level (not at all) to the next level (a few times per year), replicating the First Step Effect within a nonparametric machine learning framework. For Gen Z, this initial transition was associated with a SHAP value shift of approximately 0.15 on the log-odds scale, representing the single largest inflection point in the dependence curve. Beyond moderate socialization levels, the Gen Z SHAP dependence curve flattened markedly, mirroring the plateau observed in the logistic regression predicted probabilities. For older generations, the SHAP dependence curves maintained positive slopes at higher socialization levels, consistent with the sustained marginal returns documented in Model 1. The convergence between the parametric (logistic regression AME) and nonparametric (GBM SHAP) approaches on both the First Step Effect and the Gen Z plateau substantially strengthens the robustness of these findings, as the two methods impose fundamentally different functional form assumptions.',

        'Across the three analytic stages, a coherent and mutually reinforcing set of findings emerged. First, the socialization\u2013volunteering relationship is real, substantial, and generationally differentiated: the survey-weighted logistic regression documented a significant interaction (*p* < .001) with a distinctive plateau for Gen Z, the latent profile analysis revealed that Gen Z is concentrated in socially isolated civic profiles, and the GBM-SHAP analysis confirmed that socialization is uniquely the most important predictor of volunteering for Gen Z alone. Second, the First Step Effect, the outsized marginal gain from moving out of complete social isolation, was replicated across the regression (AME analysis), profile (the 9.0% vs. 14.6% volunteering gap between Profiles 1 and 3), and machine learning (SHAP dependence) frameworks, establishing it as the most robust individual finding of the study. Third, the temporal stability of these patterns across the pre- and post-COVID periods, documented by both the non-significant three-way interaction in Model 5 and the stable profile distributions in the LPA, supports the interpretation that Gen Z\u2019s social disconnection and its civic consequences reflect enduring cohort characteristics rather than transient pandemic effects. Together, these findings position face-to-face socialization as a pivotal and underappreciated mechanism linking generational social patterns to civic participation, with particular urgency for understanding the civic trajectory of Gen Z.',
    ]

    for para_text in results_paragraphs:
        add_body_paragraph(doc, para_text)


def _build_discussion_hardcoded(doc):
    """DEPRECATED: Build the Discussion section with hardcoded text."""
    add_heading_apa(doc, "Discussion", level=1)

    discussion_paragraphs = [
        'This study examined the association between in-person socialization frequency and volunteering across five generational cohorts using four waves of nationally representative CPS-CEV data and an integrated analytic design combining logistic regression, Latent Profile Analysis, and Gradient Boosting with TreeSHAP. The findings converge on a central conclusion: social isolation is a structural barrier to civic participation, and the youngest adult cohort is disproportionately affected.',

        'The logistic regression results (RQ1) revealed a consistent \u201cFirst Step Effect\u201d across all generations: the transition from no socialization to occasional contact (a few times per year) produced the largest marginal gain in volunteering probability, approximately 8 to 10 percentage points (*p* < .001). This finding is consistent with Cacioppo and Cacioppo\u2019s (2014) theorized cycle of social withdrawal, in which even minimal social contact disrupts the self-reinforcing dynamic of isolation and perceived social threat. Beyond this initial threshold, however, the generational trajectories diverged sharply. Generation Z\u2019s predicted volunteering probability rose from 27.3% at \u201cnot at all\u201d to approximately 48% at \u201conce per month\u201d and then plateaued, reaching only 50.8% at daily socialization. In contrast, Baby Boomers continued to gain from increased socialization, climbing from 16.7% to 46.8% across the full range. This plateau effect, unique to Gen Z, suggests that additional social contact does not translate into additional civic recruitment for this cohort in the way that social capital theory (Putnam, 2000) would predict for populations embedded in dense associational networks. Three-way interactions further showed that education significantly moderated the socialization-volunteering-generation relationship (*p* = .005), boosting the association more strongly for older cohorts, while employment was only marginally significant (*p* = .049). Notably, civic social media use did not moderate the relationship differently by generation (*p* = .964), indicating that the digital compensation effect posited by Bennett and Segerberg\u2019s (2012) connective action framework is uniform across cohorts rather than providing a generation-specific alternative pathway for digitally native young adults.',

        'Perhaps the most consequential finding concerns temporal stability. The three-way interaction between socialization, generation, and the post-COVID period was not statistically significant (*p* = .344), indicating that the socialization-volunteering-generation relationship was stable across the pre-pandemic (2017, 2019) and post-pandemic (2021, 2023) waves. This result directly addresses the concern that declining youth volunteering reflects a pandemic-era disruption rather than a deeper structural phenomenon. Gen Z\u2019s weaker conversion of social contact into volunteering was already evident before COVID-19 altered patterns of social behavior and persisted unchanged afterward. The implication is that pandemic recovery strategies alone will not resolve the civic engagement deficit among young adults, because the deficit predates the pandemic itself. This finding aligns with Twenge et al.\u2019s (2019) documentation that declines in adolescent socialization were well underway by 2017 and extends the analysis to civic consequences that Twenge\u2019s work primarily theorized rather than empirically tested.',

        'The person-centered results (RQ2) reinforced these variable-centered findings through a complementary lens. Latent Profile Analysis identified distinct civic engagement typologies, and the generational distribution across profiles was striking: 35.3% of Gen Z respondents fell into the \u201cIsolated Disengaged\u201d profile, compared with only 15.3% of Baby Boomers, while Gen Z was correspondingly underrepresented in the \u201cFully Engaged\u201d profile (3.7% versus 15.1% of Boomers). Critically, the profile distribution for Gen Z was essentially unchanged between the pre-COVID and post-COVID periods (Isolated Disengaged: 35.5% pre-COVID vs. 35.2% post-COVID), corroborating the regression finding that Gen Z\u2019s civic disconnection is structural rather than pandemic-induced.',

        'The machine learning validation (RQ3) provided a third line of convergent evidence. The GBM model achieved moderate predictive accuracy (test AUC = .731), and TreeSHAP decomposition revealed that CESOCIALIZE was the single most important predictor of volunteering for Gen Z (mean |SHAP| = 0.467), while education ranked first for all other generations. This generation-specific dominance of socialization underscores how foundational in-person contact is for Gen Z\u2019s civic participation; for older cohorts with established social networks, the traditional civic voluntarism resources of education and employment retain their predictive primacy (Verba et al., 1995). SHAP dependence plots independently recovered the First Step Effect, with the steepest gradient at the transition from the lowest to second-lowest socialization category, confirming that this threshold dynamic is not an artifact of the parametric regression specification.',

        'These findings collectively challenge the prevailing narrative about generational differences in civic engagement. The dominant interpretation, advanced by Dalton (2008) and others, frames declining youth participation in traditional civic activities as a value shift from \u201cduty-based\u201d to \u201cengaged\u201d citizenship, in which younger cohorts express civic commitment through informal and digital channels rather than through formal volunteering and organizational membership. Our results suggest a different explanation. Consider that Gen Z exhibited the highest predicted volunteering probability at the lowest socialization level: 27.3% of completely isolated Gen Z adults still volunteered, compared with only 16.7% of similarly isolated Boomers. If the decline in youth volunteering reflected weakened civic motivation, one would expect lower baseline rates among Gen Z, not higher ones. Instead, Gen Z appears to be more civically motivated at the individual level but less able to convert that motivation into sustained participation. The plateau at moderate socialization levels is more consistent with a structural account in which Gen Z lacks the dense, multiplex social networks that Granovetter (1973) identified as conduits for weak-tie recruitment. In Putnam\u2019s (2000) formulation, they are not bowling alone by choice but rather scrolling together in a digital environment that sustains the perception of connection without generating the face-to-face encounters through which civic recruitment historically occurs.',

        'These results carry three sets of practical implications. First, for volunteer recruitment and management, the First Step Effect indicates that interventions targeting the initial transition out of social isolation yield the greatest marginal return in volunteering probability. This finding favors low-threshold, single-contact strategies over sustained-commitment appeals; creating a single opportunity for in-person contact is more efficient than intensifying recruitment among already-connected individuals. For Generation Z specifically, the plateau effect suggests that digital-only recruitment campaigns are insufficient. The uniform civic social media moderation across generations (*p* = .964) means that social media use confers the same modest boost to volunteering regardless of age, offering no special compensatory advantage for the cohort most dependent on digital communication. Volunteer organizations should therefore invest in creating physical gathering spaces and initial-contact events, particularly in settings frequented by young adults such as college campuses and coworking spaces.',

        'Second, for community development and Extension programming, the findings position Cooperative Extension as a potentially critical bridge institution. Extension offices are embedded at the interface of campus and community, where socially isolated young adults might otherwise lack institutional pathways to civic participation. Klinenberg\u2019s (2018) concept of \u201csocial infrastructure\u201d, the physical spaces and institutions that facilitate social contact, provides a framework for understanding why investment in community gathering places, service-learning programs, and Extension-facilitated community events may have civic returns that exceed their direct programmatic objectives. Extension educators are uniquely positioned to design programming that treats social connection not as a byproduct of volunteering but as a prerequisite for it.',

        'Third, at the policy level, the Surgeon General\u2019s 2023 advisory on loneliness and social isolation focused predominantly on health consequences. The present findings suggest that the advisory\u2019s framing should be extended to encompass civic consequences: social isolation does not merely increase the risk of heart disease and depression but also erodes the social foundation on which democratic participation depends. Federal and state civic engagement initiatives, including AmeriCorps, national service programs, and campus-based volunteer centers, would benefit from incorporating a \u201csocial activation\u201d component that addresses social isolation as a barrier to entry rather than assuming that motivation alone is sufficient to drive participation.',

        'Several limitations should be noted. First, the pooled cross-sectional design precludes causal inference; socialization may be endogenous to volunteering if the same individuals who volunteer also seek out social contact, or if both behaviors reflect a common underlying civic disposition. Future research using true panel data or quasi-experimental designs could more definitively establish the direction of this relationship. Second, the CESOCIALIZE measure captures frequency of socialization but not its quality, depth, or content; \u201cgot together socially\u201d is a broad prompt that encompasses both the weak-tie encounters most relevant to civic recruitment and close-tie interactions that may have different civic consequences. Third, the binary COVID coding (2017/2019 vs. 2021/2023) is approximate; the 2021 wave was administered during, not after, the pandemic, and the pre/post distinction does not capture the heterogeneity of pandemic experiences across regions and communities. Fourth, generational classification by birth cohort inherently conflates age, period, and cohort effects; while the four-wave design provides some leverage for disentangling these factors, a formal age-period-cohort decomposition would require additional data waves. Fifth, the PES16F variable measuring virtual versus in-person volunteering mode is available only for volunteers in the 2023 wave, preventing systematic analysis of how volunteering modality interacts with social isolation across generations. Sixth, all measures are self-reported and subject to social desirability bias, which may differentially affect reporting of socially valued behaviors such as volunteering.',
    ]

    for para_text in discussion_paragraphs:
        add_body_paragraph(doc, para_text)

    # Conclusion as Heading 2
    add_heading_apa(doc, "Conclusion", level=2)

    conclusion_text = (
        'In conclusion, the loneliness epidemic is also a civic engagement crisis. '
        'This study demonstrates that in-person socialization is the strongest predictor '
        'of volunteering for the most socially isolated generation in American survey history, '
        'and that the initial step out of isolation, from no social contact to occasional '
        'contact, produces the largest marginal civic return. This pattern is stable across '
        'four survey waves spanning the COVID-19 pandemic, confirming a structural cohort '
        'effect rather than a pandemic artifact. For Generation Z, the path to volunteering '
        'begins not with appeals to duty or digital campaigns, but with creating the physical '
        'social encounters that make civic recruitment possible. Addressing the decline in '
        'youth volunteering therefore requires not only rethinking recruitment strategies but '
        'also rebuilding the social infrastructure through which civic participation has '
        'historically been sustained.'
    )
    add_body_paragraph(doc, conclusion_text)


def build_references(doc):
    """Build the References section."""
    insert_page_break(doc)
    add_heading_apa(doc, "References", level=1)

    refs = [
        "AmeriCorps. (2023). Volunteering and civic life in America.",
        "Arel-Bundock, V. (2023). *marginaleffects: Predictions, comparisons, slopes, marginal means, and hypothesis tests*. R package.",
        "Bennett, W. L., & Segerberg, A. (2012). The logic of connective action. *Information, Communication & Society, 15*(5), 739\u2013768.",
        "Beyerlein, K., & Hipp, J. R. (2006). From pews to participation: The effect of congregation activity and context on bridging civic engagement. *Social Problems, 53*(1), 97\u2013117.",
        "Bode, L., Vraga, E. K., Borah, P., & Shah, D. V. (2014). A new space for political behavior. *Political Communication, 31*(1), 52\u201374.",
        "Brady, H. E., Verba, S., & Schlozman, K. L. (1995). Beyond SES: A resource model of political participation. *American Political Science Review, 89*(2), 271\u2013294.",
        "Bureau of Labor Statistics. (2016). *Volunteering in the United States, 2015*.",
        "Cacioppo, J. T., & Cacioppo, S. (2014). Social relationships and health: The toxic effects of perceived social isolation. *Social and Personality Psychology Compass, 8*(2), 58\u201372.",
        "Coleman, J. S. (1988). Social capital in the creation of human capital. *American Journal of Sociology, 94*, S95\u2013S120.",
        "Collins, L. M., & Lanza, S. T. (2010). *Latent class and latent transition analysis*. Wiley.",
        "Corporation for National and Community Service. (2018). *Volunteering in America*.",
        "Dalton, R. J. (2008). *The good citizen: How a younger generation is reshaping American politics*. CQ Press.",
        "Flanagan, C. A., & Levine, P. (2010). Civic engagement and the transition to adulthood. *The Future of Children, 20*(1), 159\u2013179.",
        "Flood, S., King, M., Rodgers, R., Ruggles, S., Warren, J. R., & Westberry, M. (2023). *Integrated Public Use Microdata Series, Current Population Survey: Version 11.0*. IPUMS.",
        "Friedman, J. H. (2001). Greedy function approximation: A gradient boosting machine. *Annals of Statistics, 29*(5), 1189\u20131232.",
        "Gil de Zuniga, H., Jung, N., & Valenzuela, S. (2012). Social media use for news and individuals\u2019 social capital, civic engagement and political participation. *Journal of Computer-Mediated Communication, 17*(3), 319\u2013336.",
        "Granovetter, M. S. (1973). The strength of weak ties. *American Journal of Sociology, 78*(6), 1360\u20131380.",
        "Holt-Lunstad, J., Smith, T. B., & Layton, J. B. (2010). Social relationships and mortality risk: A meta-analytic review. *PLoS Medicine, 7*(7), e1000316.",
        "Hooghe, M. (2012). Taking part in politics? An analysis of macro-level determinants and individual-level explanations. *International Review of Sociology, 22*(1), 9\u201326.",
        "Jennings, M. K., & Stoker, L. (2004). Social trust and civic engagement across time and generations. *Acta Politica, 39*, 342\u2013379.",
        "Kelle, N., Simonson, J., & Gordo, L. R. (2025). Generational differences in voluntary engagement among older adults. *Voluntas, 36*, 1\u201314.",
        "Klinenberg, E. (2018). *Palaces for the people: How social infrastructure can help fight inequality, polarization, and the decline of civic life*. Crown.",
        "Lim, C., & Laurence, J. (2015). Doing good when times are bad: Volunteering behaviour in economic hard times. *British Journal of Sociology, 66*(2), 319\u2013344.",
        "Lundberg, S. M., Erion, G., Chen, H., DeGrave, A., Prutkin, J. M., Nair, B., ... & Lee, S. I. (2020). From local explanations to global understanding with explainable AI for trees. *Nature Machine Intelligence, 2*(1), 56\u201367.",
        "Lundberg, S. M., & Lee, S. I. (2017). A unified approach to interpreting model predictions. *Advances in Neural Information Processing Systems, 30*.",
        "Manturuk, K., Lindblad, M., & Quercia, R. (2012). Homeownership and civic engagement in low-income urban neighborhoods. *Urban Affairs Review, 48*(5), 731\u2013760.",
        "Mauss, M. (1925/2000). *The gift: The form and reason for exchange in archaic societies*. W. W. Norton.",
        "McPherson, M., Smith-Lovin, L., & Brashears, M. E. (2006). Social isolation in America: Changes in core discussion networks over two decades. *American Sociological Review, 71*(3), 353\u2013375.",
        "Musick, M. A., & Wilson, J. (2008). *Volunteers: A social profile*. Indiana University Press.",
        "Muth\u00e9n, B. O., & Muth\u00e9n, L. K. (2000). Integrating person-centered and variable-centered analyses. *Alcoholism: Clinical and Experimental Research, 24*(5), 543\u2013550.",
        "Nesbit, R. (2012). The influence of major life cycle events on volunteering. *Nonprofit and Voluntary Sector Quarterly, 41*(6), 1153\u20131174.",
        "Norris, P. (2002). *Democratic phoenix: Reinventing political activism*. Cambridge University Press.",
        "Office of the Surgeon General. (2023). *Our epidemic of loneliness and isolation: The U.S. Surgeon General\u2019s advisory on the healing effects of social connection and community*.",
        "Putnam, R. D. (2000). *Bowling alone: The collapse and revival of American community*. Simon & Schuster.",
        "Rotolo, T., & Wilson, J. (2006). Employment sector and volunteering: The contribution of nonprofit and public sector workers to the volunteer labor force. *Sociological Quarterly, 47*(1), 21\u201340.",
        "Schlozman, K. L., Verba, S., & Brady, H. E. (2012). *The unheavenly chorus: Unequal political voice and the broken promise of American democracy*. Princeton University Press.",
        "Snyder, M., & Omoto, A. M. (2008). Volunteerism: Social issues perspectives and social policy implications. *Social Issues and Policy Review, 2*(1), 1\u201336.",
        "Stukas, A. A., Hoye, R., Nicholson, M., Brown, K. M., & Aisbett, L. (2016). Motivations to volunteer and their associations with volunteers\u2019 well-being. *Nonprofit and Voluntary Sector Quarterly, 45*(1), 112\u2013132.",
        "Taniguchi, H. (2012). The determinants of formal and informal volunteering. *Voluntas, 23*, 920\u2013939.",
        "Tocqueville, A. de. (1835/2000). *Democracy in America*. University of Chicago Press.",
        "Twenge, J. M. (2017). *iGen: Why today\u2019s super-connected kids are growing up less rebellious, more tolerant, less happy, and completely unprepared for adulthood*. Atria Books.",
        "Twenge, J. M., & Park, H. (2019). The decline in adult activities among U.S. adolescents, 1976\u20132016. *Child Development, 90*(2), 638\u2013654.",
        "Twenge, J. M., Spitzberg, B. H., & Campbell, W. K. (2019). Less in-person social interaction with peers among U.S. adolescents in the 21st century and links to loneliness. *Journal of Social and Personal Relationships, 36*(6), 1892\u20131913.",
        "Verba, S., Schlozman, K. L., & Brady, H. E. (1995). *Voice and equality: Civic voluntarism in American politics*. Harvard University Press.",
        "Vermunt, J. K., & Magidson, J. (2002). Latent class cluster analysis. In J. A. Hagenaars & A. L. McCutcheon (Eds.), *Applied latent class analysis* (pp. 89\u2013106). Cambridge University Press.",
        "Wilson, J. (2000). Volunteering. *Annual Review of Sociology, 26*, 215\u2013240.",
        "Wilson, J. (2012). Volunteerism research: A review essay. *Nonprofit and Voluntary Sector Quarterly, 41*(2), 176\u2013212.",
        "Wray-Lake, L., & Hart, D. (2012). Growing social inequalities in youth civic engagement? Evidence from the National Election Study. *PS: Political Science & Politics, 45*(3), 456\u2013461.",
        "Zukin, C., Keeter, S., Andolina, M., Jenkins, K., & Delli Carpini, M. X. (2006). *A new engagement? Political participation, civic life, and the changing American citizen*. Oxford University Press.",
    ]

    for ref in refs:
        add_reference(doc, ref)


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    import os

    doc = setup_document()

    # Build all sections in order
    build_title_page(doc)
    build_abstract(doc)

    # All body sections read from markdown source files
    paper_dir = os.path.dirname(os.path.abspath(__file__))
    for md_file in [
        "01_introduction.md",
        "02_theoretical_framework.md",
        "02a_purpose_rq.md",
        "03_method.md",
        "04_results.md",
        "05_discussion.md",
    ]:
        build_section_from_md(doc, os.path.join(paper_dir, md_file))

    build_references(doc)

    # Save
    output_path = os.path.join(
        os.path.dirname(os.path.abspath(__file__)),
        "Bowling_Alone_CLEAN.docx"
    )
    doc.save(output_path)
    print(f"Document saved to: {output_path}")
    print("Done.")


if __name__ == "__main__":
    main()
